import os
import sys
import json
import webbrowser
import threading
import time
import datetime
import traceback
import struct
import csv
import math
import subprocess
import smtplib
import shutil
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from flask import Flask, render_template_string, request, jsonify, send_file, session, redirect, url_for
from waitress import serve



import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg') 
from matplotlib.figure import Figure
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import matplotlib.ticker as ticker
import matplotlib.gridspec as gridspec
from mpl_toolkits.mplot3d import Axes3D 

try:
    from docx import Document
    from docx.shared import Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False

# ==============================================================================
# CONFIGURACIÓN DEL SERVIDOR Y SEGURIDAD
# ==============================================================================
ARCHIVO_FUENTES_GLOBAL = "fuentes_globales.json"
ARCHIVO_USUARIOS = "usuarios.json"
PUERTO = 8082
OFFSET_ALERTAS = 15.0
MAX_PICOS = 30

app = Flask(__name__)
app.secret_key = "phoenix_argus_super_secret_key_2026" 

# ==============================================================================
# GESTORES AUXILIARES (RED, USUARIOS Y EMAIL)
# ==============================================================================
def conectar_red_windows(ruta, usr, pwd):
    """Utiliza el formato exacto original del Bot_Analisis_UMA para conectar"""
    if not usr or not pwd or not str(ruta).startswith('\\\\'): return
    if os.path.exists(ruta): return 

    # Extraer carpeta base para que net use funcione sin colgarse con archivos
    ruta_base = ruta
    nombre_archivo = os.path.basename(ruta)
    if '.' in nombre_archivo and len(nombre_archivo.split('.')[-1]) <= 4:
        ruta_base = os.path.dirname(ruta)

    # Comando original que funcionaba perfectamente:
    comando = f'net use "{ruta_base}" {pwd} /user:{usr}'
    try: 
        subprocess.run(comando, shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=5)
    except: pass

def cargar_usuarios():
    if os.path.exists(ARCHIVO_USUARIOS):
        try:
            with open(ARCHIVO_USUARIOS, 'r', encoding='utf-8') as f: return json.load(f)
        except: pass
    default_users = {"admin": {"password": "Argus2026", "role": "admin"}}
    with open(ARCHIVO_USUARIOS, 'w', encoding='utf-8') as f: json.dump(default_users, f, indent=4)
    return default_users

def enviar_correo_reporte(destinatarios_str, ruta_salida, prefijo, titulo, hay_alertas):
    if not destinatarios_str.strip(): return
    destinatarios = [e.strip() for e in destinatarios_str.split(',') if e.strip()]
    SMTP_SERVER = "smtp.gmail.com" 
    SMTP_PORT = 587 
    SMTP_USER = "jpitmalagaalertas@gmail.com" 
    SMTP_PASS = "kwejengkvwuahmim"
    EMAIL_ORIGEN = "jpitmalagaalertas@gmail.com" 
    
    msg = MIMEMultipart()
    msg['From'] = EMAIL_ORIGEN
    msg['To'] = ", ".join(destinatarios)
    msg['Subject'] = f"📊 Reporte ARGUS - {prefijo} ({titulo})"
    
    body_text = f"Se ha completado la extracción de datos solicitada desde el Dashboard para {prefijo}.\n\n"
    if hay_alertas:
        body_text += "⚠️ ATENCIÓN: Se han detectado picos de alerta en los tramos analizados.\n"
        body_text += "Se adjunta el reporte detallado con las gráficas de las incidencias."
    else:
        body_text += "✅ Todo correcto: No se han detectado alertas (no hay picos por encima del umbral) en los tramos solicitados.\n"
        body_text += "Se adjunta el reporte vacío para su constancia."

    msg.attach(MIMEText(body_text, 'plain'))
    
    total_size_bytes = 0
    MAX_EMAIL_SIZE_BYTES = 20 * 1024 * 1024 # 20MB límite prudencial para Gmail
    archivos_omitidos = False
    
    for root, _, files in os.walk(ruta_salida):
        for file in files:
            # El usuario ha pedido explícitamente SOLO recibir el reporte de alertas (Word)
            if not file.endswith('.docx'):
                continue
                
            ruta_arch = os.path.join(root, file)
            size = os.path.getsize(ruta_arch)
            
            if (total_size_bytes + size) > MAX_EMAIL_SIZE_BYTES:
                archivos_omitidos = True
                continue
                
            try:
                with open(ruta_arch, "rb") as f:
                    part = MIMEApplication(f.read(), Name=file)
                    part['Content-Disposition'] = f'attachment; filename="{file}"'
                    msg.attach(part)
                    total_size_bytes += size
            except: pass
            
    if archivos_omitidos:
        msg.attach(MIMEText("\n\n⚠️ NOTA DEL SISTEMA: Algunos reportes no se adjuntaron porque excedían el límite de 20MB acordado para el correo. Revisa la carpeta en red compartida.", 'plain'))
        
    try:
        # Timeout extendido a 60s
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT, timeout=60)
        server.starttls()
        if SMTP_USER: server.login(SMTP_USER, SMTP_PASS)
        server.sendmail(EMAIL_ORIGEN, destinatarios, msg.as_string())
        server.quit()
        print(f"📧 Email enviado a {destinatarios} ({total_size_bytes/1024/1024:.1f} MB adjuntos)")
    except Exception as e: print(f"❌ Error enviando correo: {e}")
    sys.stdout.flush()

# ==============================================================================
# MIDDLEWARE DE SEGURIDAD
# ==============================================================================
@app.before_request
def requerir_login():
    if request.endpoint not in ['login', 'static'] and not session.get('autenticado'):
        if request.path.startswith('/api/'): return jsonify({"error": "Acceso denegado. Requiere autenticación."}), 401
        return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    error = ""
    if request.method == 'POST':
        u = request.form.get('username')
        p = request.form.get('password')
        db = cargar_usuarios()
        if u in db and db[u]['password'] == p:
            session['autenticado'] = True; session['usuario_actual'] = u; session['rol'] = db[u]['role']
            return redirect(url_for('home'))
        else: error = "Usuario o contraseña incorrectos."
            
    return f"""
    <!DOCTYPE html>
    <html lang="es">
    <head><meta charset="UTF-8"><title>ARGUS - Acceso</title>
    <link rel="icon" type="image/x-icon" href="/favicon.ico">
    <style>body {{ background: #0b0f19; color: white; font-family: sans-serif; display: flex; justify-content: center; align-items: center; height: 100vh; margin: 0; }} .login-box {{ background: #161b22; padding: 40px; border-radius: 12px; border-top: 4px solid #3b82f6; text-align: center; box-shadow: 0 10px 30px rgba(0,0,0,0.5); width: 300px; }} input {{ width: 100%; padding: 12px; margin: 10px 0; border-radius: 6px; border: 1px solid #30363d; background: #0b0f19; color: white; box-sizing: border-box; outline: none; }} button {{ width: 100%; padding: 12px; background: #3b82f6; color: white; border: none; border-radius: 6px; font-weight: bold; cursor: pointer; margin-top: 10px;}} button:hover {{ background: #2563eb; }}</style></head>
    <body><div class="login-box"><h1 style="margin:0;">🦅 ARGUS</h1><p style="color:#8b949e; font-size:14px; margin-bottom:25px;">Centro de Control Argus</p><form method="POST"><input type="text" name="username" placeholder="Usuario" required autocomplete="off" autofocus><input type="password" name="password" placeholder="Contraseña" required><button type="submit">ENTRAR</button></form><div style="color:#ef4444; margin-top:15px; font-size:13px;">{error}</div></div></body></html>
    """

@app.route('/logout')
def logout(): session.clear(); return redirect(url_for('login'))

# ==============================================================================
# MOTOR MATEMÁTICO Y GRÁFICO
# ==============================================================================
def calcular_metricas(df, offset_db, max_picos=30):
    p5_ruido = np.percentile(df['L'], 5)
    umbral = p5_ruido + offset_db
    df_maxhold = df.groupby(df['F'].round(4))['L'].max().reset_index()
    
    # Ordenar por tiempo para detección cronológica
    df_over = df[df['L'] >= umbral].sort_values(by='T', ascending=True)
    picos_detectados = []
    
    for _, row in df_over.iterrows():
        es_nuevo = True
        for p in picos_detectados:
            if abs((row['T'] - p['T']).total_seconds()) < 60 and abs(row['F'] - p['F']) < 0.5:
                es_nuevo = False
                # Capturar el APEX (valor más alto) dentro de este evento cronológico
                if row['L'] > p['L']:
                    p['T'] = row['T']; p['L'] = row['L']; p['F'] = row['F']
                break
        if es_nuevo:
            if len(picos_detectados) < max_picos:
                picos_detectados.append({'T': row['T'], 'F': row['F'], 'L': row['L']})
                
    df_picos = pd.DataFrame(picos_detectados)
    if not df_picos.empty:
        df_picos = df_picos.sort_values('T').reset_index(drop=True)
        df_picos['Numero'] = df_picos.index + 1 
        return df_maxhold, df_picos, p5_ruido, umbral, df_picos['F'].min(), df_picos['F'].max(), (df_picos['F'].min() + df_picos['F'].max()) / 2
    return df_maxhold, df_picos, p5_ruido, umbral, 0.0, 0.0, 0.0

def dibujar_panel_lateral(ax_side, umbral, df_picos, offset_db):
    ax_side.axis('off'); ax_side.set_xlim(0, 1); ax_side.set_ylim(0, 1)
    ax_side.text(0.0, 0.98, "LEYENDA", fontsize=11, fontweight='bold', ha='left', va='top')
    ax_side.plot([0.0, 0.15], [0.93, 0.93], color='#0055A4', lw=1.5)
    ax_side.text(0.2, 0.93, "Nivel de Señal", fontsize=10, va='center')
    ax_side.plot([0.0, 0.15], [0.88, 0.88], color='orange', ls='--', lw=1.5)
    ax_side.text(0.2, 0.88, f"Umbral (+{offset_db}dB)", fontsize=10, va='center')
    ax_side.scatter([0.075], [0.83], color='red', s=40, zorder=5)
    ax_side.text(0.2, 0.83, "Detecciones", fontsize=10, va='center')
    ax_side.text(0.0, 0.75, f"TOP PICOS ({len(df_picos)})", fontsize=11, fontweight='bold', ha='left', va='top')
    
    if df_picos.empty: ax_side.text(0.0, 0.70, "Ningún pico detectado.", fontsize=9, color='gray')
    else:
        ax_side.text(0.0, 0.70, "Nº   Hora       Frec(MHz)   Nivel", fontsize=9, fontweight='bold', fontfamily='monospace')
        y_pos = 0.67; step = min(0.03, 0.65 / max(15, len(df_picos)))
        for _, row in df_picos.iterrows():
            ax_side.text(0.0, y_pos, f"{int(row['Numero']):02d} | {row['T'].strftime('%H:%M:%S')} | {row['F']:8.4f} | {row['L']:5.1f}", fontsize=8, fontfamily='monospace', va='center')
            y_pos -= step

# ==============================================================================
# FRONTEND: HTML / CSS / JS
# ==============================================================================
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>ARGUS - Panel Web Unificado</title>
    <link rel="icon" type="image/x-icon" href="/favicon.ico">
    <link rel="stylesheet" href="https://cdn.datatables.net/1.13.6/css/jquery.dataTables.min.css">
    <script src="https://code.jquery.com/jquery-3.7.0.min.js"></script>
    <script src="https://cdn.datatables.net/1.13.6/js/jquery.dataTables.min.js"></script>

    <style>
        :root { --bg: #0b0f19; --panel: #161b22; --panel-light: #21262d; --text: #e2e8f0; --accent: #3b82f6; --accent-hover: #2563eb; --danger: #ef4444; --success: #10b981; --warning: #f59e0b; }
        html, body { height: 100%; margin: 0; padding: 0; overflow: hidden; background-color: var(--bg); color: var(--text); font-family: 'Segoe UI', system-ui, sans-serif; }
        body { padding: 15px; box-sizing: border-box; display: flex; justify-content: center; }
        
        .dataTables_wrapper .dataTables_length, .dataTables_wrapper .dataTables_filter, .dataTables_wrapper .dataTables_info, .dataTables_wrapper .dataTables_processing, .dataTables_wrapper .dataTables_paginate { color: #8b949e !important; font-size: 13px; margin-bottom: 10px; }
        .dataTables_wrapper .dataTables_filter input { background-color: var(--bg); color: white; border: 1px solid var(--panel-light); border-radius: 4px; padding: 5px; margin-left: 10px; }
        table.dataTable tbody tr { background-color: var(--bg) !important; color: white !important; }
        table.dataTable tbody tr:hover { background-color: var(--panel-light) !important; }
        table.dataTable thead th, table.dataTable tfoot th { color: #8b949e; border-bottom: 1px solid var(--panel-light) !important; background-color: #1a1f26; }
        table.dataTable { border-collapse: collapse !important; }
        table.dataTable.no-footer { border-bottom: 1px solid var(--panel-light); }

        .main-wrapper { width: 100%; max-width: 1600px; height: 100%; display: flex; flex-direction: column; gap: 15px; }
        .header-panel { background-color: var(--panel); padding: 15px 30px; border-radius: 12px; border-top: 4px solid var(--accent); display: flex; justify-content: space-between; align-items: center; flex-shrink: 0;}
        h1 { margin: 0; font-size: 24px; color: #fff; letter-spacing: 1px; } .subtitle { color: #8b949e; font-size: 13px; margin-top: 5px;}
        .tabs-container { display: flex; gap: 10px; background: var(--panel); padding: 8px; border-radius: 10px; flex-shrink: 0;}
        .tab-btn { background: transparent; color: #8b949e; border: none; padding: 10px 20px; font-size: 14px; font-weight: bold; border-radius: 6px; cursor: pointer; transition: 0.3s; flex: 1;}
        .tab-btn.active { background: var(--accent); color: white; } .tab-btn:hover:not(.active) { background: var(--panel-light); color: white; }
        .tab-content { flex: 1; display: flex; flex-direction: column; gap: 15px; min-height: 0; overflow-y: auto; padding-right: 5px;}
        .tab-content::-webkit-scrollbar { width: 8px; } .tab-content::-webkit-scrollbar-thumb { background: #30363d; border-radius: 10px; }
        
        .grid-2col { display: grid; grid-template-columns: 1fr 1fr; gap: 20px; } .grid-3col { display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 20px; }
        .panel { background-color: var(--panel); padding: 20px; border-radius: 12px; border: 1px solid #21262d; display: flex; flex-direction: column; position: relative;}
        .panel-title { font-size: 15px; font-weight: 600; color: #fff; margin-top: 0; margin-bottom: 15px; border-bottom: 1px solid var(--panel-light); padding-bottom: 10px; display: flex; justify-content: space-between; align-items: center;}
        
        .card-link { cursor: pointer; transition: transform 0.2s; } .card-link:hover { transform: translateY(-3px); border-color: rgba(59, 130, 246, 0.4); }
        input[type="text"], input[type="date"], input[type="time"], select, textarea, input[type="password"] { background-color: var(--bg); border: 1px solid var(--panel-light); color: white; padding: 10px; border-radius: 6px; font-size: 14px; outline: none; width: 100%; box-sizing: border-box; margin-bottom: 15px;}
        input:focus, select:focus, textarea:focus { border-color: var(--accent); }
        .input-group { display: flex; gap: 15px; } .input-group > div { flex: 1; }
        label { display: block; font-size: 11px; color: #8b949e; font-weight: bold; margin-bottom: 5px; text-transform: uppercase;}
        
        .btn { padding: 10px 15px; border-radius: 6px; font-weight: bold; border: none; cursor: pointer; transition: 0.2s; font-size: 13px; text-decoration:none; display: inline-block;}
        .btn-primary { background-color: var(--accent); color: white; } .btn-primary:hover { background-color: var(--accent-hover); }
        .btn-danger { background-color: var(--danger); color: white; } 
        .btn-success { background-color: var(--success); color: white; width: 100%; padding: 15px; font-size: 15px; margin-top: 10px;}
        .btn-warning { background-color: var(--warning); color: black; }
        .btn-outline { background: transparent; border: 1px solid var(--panel-light); color: #fff;} .btn-outline:hover { background: var(--panel-light); }
        .check-container { display: flex; align-items: center; gap: 10px; margin-bottom: 10px; cursor: pointer; } .check-container input { width: 16px; height: 16px; accent-color: var(--accent); cursor: pointer; margin: 0;} .check-label { font-size: 13px; color: #c9d1d9; }
        
        #tab-visor { overflow: hidden; } 
        .visor-container { display: flex; gap: 15px; flex: 1; min-height: 0; }
        .col-explorer { width: 320px; min-width: 320px; max-width: 320px; display: flex; flex-direction: column; flex-shrink: 0; border: 1px solid var(--panel-light); }
        .col-preview { flex: 1; display: flex; flex-direction: column; min-width: 0; border: 1px solid var(--panel-light); overflow: hidden;}
        #visor-lista { flex: 1; overflow-y: auto; overflow-x: hidden; padding-right: 5px; }
        #visor-contenido { flex: 1; overflow: auto; background: var(--bg); border-radius: 8px; position: relative; display: flex; justify-content: center; align-items: flex-start;}
        .file-item { padding: 8px; cursor: pointer; border-radius: 6px; display: flex; align-items: center; gap: 10px; color: #c9d1d9; font-size: 13px; border-bottom: 1px solid #1a1f26;}
        .file-item:hover { background: var(--panel-light); color: white; }
        .file-item.folder { font-weight: bold; color: var(--accent); }
        .file-item.active { background: #1a273d; border-left: 3px solid var(--accent); }

        .zoom-controls { display: none; background: #1a1f26; padding: 5px 15px; border-radius: 20px; border: 1px solid #30363d; align-items: center; gap: 15px; margin-right: 15px;}
        .zoom-btn { background: transparent; color: white; border: none; font-size: 18px; cursor: pointer; font-weight: bold; padding: 0 5px;}
        .zoom-btn:hover { color: var(--accent); }
        #docx-container { transform-origin: top center; transition: transform 0.2s ease-out; }

        .nav-img-btn { position: absolute; top: 50%; transform: translateY(-50%); background: rgba(0,0,0,0.6); color: white; border: none; font-size: 24px; padding: 15px 10px; cursor: pointer; z-index: 50; border-radius: 8px; transition: 0.2s; opacity: 0.3;}
        .nav-img-btn:hover { background: rgba(59, 130, 246, 0.8); opacity: 1;}
        .nav-left { left: 10px; } .nav-right { right: 10px; }

        #overlay-loading { display:none; position:fixed; top:0; left:0; width:100%; height:100%; background:rgba(0,0,0,0.8); z-index:2000; color:white; justify-content:center; align-items:center; flex-direction:column;}
        .loader { border: 6px solid #161b22; border-top: 6px solid var(--accent); border-radius: 50%; width: 50px; height: 50px; animation: spin 1s linear infinite; margin-bottom: 20px;}
        @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
        
        .switch { position: relative; display: inline-block; width: 50px; height: 24px; } .switch input { opacity: 0; width: 0; height: 0; }
        .slider { position: absolute; cursor: pointer; top: 0; left: 0; right: 0; bottom: 0; background-color: #30363d; transition: .4s; border-radius: 34px; } .slider:before { position: absolute; content: ""; height: 16px; width: 16px; left: 4px; bottom: 4px; background-color: white; transition: .4s; border-radius: 50%; }
        input:checked + .slider { background-color: var(--success); } input:checked + .slider:before { transform: translateX(26px); }
        .modal-overlay { display: none; position: fixed; top: 0; left: 0; width: 100%; height: 100%; background: rgba(0,0,0,0.7); z-index: 1000; justify-content: center; align-items: center; } .modal { background: var(--panel); width: 650px; padding: 25px; border-radius: 12px; }
        .fuente-item { background: var(--bg); padding: 15px; border-radius: 8px; margin-bottom: 10px; border-left: 4px solid var(--accent); display: flex; justify-content: space-between; align-items: center;}
        
        .calendar-grid { display: grid; grid-template-columns: repeat(7, 1fr); gap: 5px; text-align: center; }
        .day-name { font-size: 11px; font-weight: bold; color: #8b949e; padding: 5px 0; }
        .day { padding: 8px 0; border-radius: 6px; cursor: pointer; font-size: 13px; font-weight: bold; margin-bottom: 2px; }
        .day.active { background: var(--panel-light); color: white; transition: 0.2s;} .day.active:hover { background: var(--accent); }
        .day.empty { background: transparent; cursor: default; }
        .day.blocked { background: rgba(239, 68, 68, 0.2); color: var(--danger); border: 1px solid rgba(239, 68, 68, 0.4); transition: 0.2s;} 
        .day.blocked:hover { background: rgba(239, 68, 68, 0.4); }
    </style>
</head>
<body>
    <div id="overlay-loading">
        <div class="loader"></div>
        <h2 id="loading-text" style="text-align: center;">Procesando...<br><span style="font-size: 14px; font-weight: normal; color: #8b949e;">Mira la consola CMD para detalles en vivo.</span></h2>
    </div>

    <div class="main-wrapper">
        <div class="header-panel">
            <div>
                <h1>🦅 ARGUS <span style="font-weight: 300; opacity: 0.8;">| Centro Analítico Web</span></h1>
                <div class="subtitle">Extracción a Demanda y Control de Alertas Automáticas</div>
            </div>
            <div style="display: flex; align-items: center; gap: 15px;">
                {% if es_admin %}
                <button type="button" class="btn btn-outline" style="border-color: var(--warning); color: var(--warning);" onclick="abrirModalUsuarios()">👥 Gestión Usuarios</button>
                {% endif %}
                <a href="/logout" class="btn btn-outline" style="color:#ef4444; border-color:#ef4444; text-decoration:none;">Cerrar Sesión</a>
            </div>
        </div>

        <div class="tabs-container">
            <button type="button" class="tab-btn active" onclick="cambiarPestana('tab-home', this)">🏠 Inicio</button>
            <button type="button" id="btn-tab-extraer" class="tab-btn" onclick="cambiarPestana('tab-extraer', this)">📊 Extraer Resultados</button>
            <button type="button" id="btn-tab-alertas" class="tab-btn" onclick="cambiarPestana('tab-alertas', this)">⏰ Configurar Alertas</button>
            <button type="button" id="btn-tab-visor" class="tab-btn" onclick="cambiarPestana('tab-visor', this)">📂 Visor de Archivos</button>
        </div>

        <div id="tab-home" class="tab-content" style="display: flex;">
            <div class="grid-3col">
                <div class="panel card-link" style="align-items: center; text-align: center;" onclick="document.getElementById('btn-tab-extraer').click()">
                    <h1 style="font-size: 40px; margin-bottom: 10px;">📊</h1><h3>Extracción Directa</h3><p style="color: #8b949e; font-size: 13px;">Genera gráficas y reportes a demanda.</p>
                </div>
                <div class="panel card-link" style="align-items: center; text-align: center;" onclick="document.getElementById('btn-tab-alertas').click()">
                    <h1 style="font-size: 40px; margin-bottom: 10px;">⏰</h1><h3>Gestión de Alertas</h3><p style="color: #8b949e; font-size: 13px;">Modifica JSON, horarios y destinatarios.</p>
                </div>
                <div class="panel card-link" style="align-items: center; text-align: center;" onclick="document.getElementById('btn-tab-visor').click()">
                    <h1 style="font-size: 40px; margin-bottom: 10px;">📂</h1><h3>Visor de Archivos</h3><p style="color: #8b949e; font-size: 13px;">Navega, visualiza y descarga reportes PDF, CSV o imágenes.</p>
                </div>
            </div>
        </div>

        <div id="tab-extraer" class="tab-content" style="display: none;">
            <div class="grid-2col">
                <div class="panel">
                    <div class="panel-title"><span>1. Origen y Fechas</span><button type="button" class="btn btn-outline" style="padding: 5px 10px; font-size: 11px;" onclick="abrirModalFuentes()">⚙️ Configurar Fuentes</button></div>
                    <label>Seleccione la Fuente de Datos</label>
                    <select id="ext_fuente" onchange="actualizarPlaceholderSalida()"></select>
                    
                    <label>Tipo de Extracción</label>
                    <div style="background: var(--bg); padding: 12px; border-radius: 8px; margin-bottom: 15px; border: 1px solid var(--panel-light);">
                        <label class="check-container"><input type="radio" name="origen_datos" value="csv" checked><span class="check-label" style="color:var(--accent); font-weight:bold;">CSVs de la Carpeta Resultados (Rápido)</span></label>
                        <label class="check-container"><input type="radio" name="origen_datos" value="binario"><span class="check-label">Binarios Originales (Proceso Completo)</span></label>
                    </div>

                    <div class="input-group">
                        <div><label>Fecha Inicio</label><input type="date" id="ext_f_ini"><input type="time" id="ext_h_ini" value="00:00"></div>
                        <div><label>Fecha Fin</label><input type="date" id="ext_f_fin"><input type="time" id="ext_h_fin" value="23:59"></div>
                    </div>
                    <label>Carpeta de Salida Manual (Opcional)</label>
                    <div style="display:flex; gap:10px; margin-bottom: 15px;">
                        <input type="text" id="ext_salida" placeholder="Si está vacío, se creará: ejecucion_a_peticion_FECHA" style="margin-bottom:0; flex:1;">
                        <button type="button" class="btn btn-outline" style="height: 42px;" onclick="abrirWebBrowser('ext_salida', 'folder')">🔍 Buscar en Servidor</button>
                    </div>
                </div>

                <div class="panel">
                    <div class="panel-title">2. Entregables y Avisos</div>
                    <div class="grid-2col" style="margin-bottom: 15px;">
                        <div>
                            <label class="check-container"><input type="checkbox" id="chk_2dt" checked><span class="check-label">Gráfica Tiempo 2D</span></label>
                            <label class="check-container"><input type="checkbox" id="chk_2df" checked><span class="check-label">Gráfica Frecuencia 2D</span></label>
                            <label class="check-container"><input type="checkbox" id="chk_3d"><span class="check-label">Gráfica 3D</span></label>
                        </div>
                        <div>
                            <label class="check-container"><input type="checkbox" id="chk_csv" checked><span class="check-label">Archivo CSV</span></label>
                            <label class="check-container"><input type="checkbox" id="chk_excel" checked><span class="check-label">Archivo Excel</span></label>
                            <label class="check-container"><input type="checkbox" id="chk_word"><span class="check-label">Reporte Word Alertas</span></label>
                        </div>
                    </div>
                    
                    <div style="background: rgba(59, 130, 246, 0.1); padding: 15px; border-radius: 8px; border: 1px solid rgba(59, 130, 246, 0.3); margin-top: 15px;">
                        <label class="check-container" style="margin-bottom: 5px;">
                            <input type="checkbox" id="chk_email" onchange="document.getElementById('email_box').style.display = this.checked ? 'block' : 'none'">
                            <span class="check-label" style="font-weight: bold; color: var(--accent);">Enviar por Email al terminar</span>
                        </label>
                        <div id="email_box" style="display: none; margin-top: 10px;">
                            <label>Instrucciones: Separa los correos con coma (,)</label>
                            <textarea id="ext_emails" rows="2" placeholder="ejemplo@dsic.es, jefe@dsic.es"></textarea>
                        </div>
                    </div>

                    <div style="background: var(--bg); padding: 12px; border-radius: 8px; margin-top: 15px; border: 1px solid var(--panel-light);">
                        <label class="check-container" style="margin-bottom: 5px;">
                            <input type="checkbox" id="chk_dividir" checked onchange="document.getElementById('div_horas_box').style.display = this.checked ? 'flex' : 'none'">
                            <span class="check-label" style="font-weight: bold; color: white;">Segmentar gráficas por tramos de horas</span>
                        </label>
                        <div id="div_horas_box" style="display: flex; gap: 10px; align-items: center; margin-top: 5px;">
                            <span style="font-size: 13px; color: #8b949e;">Duración del tramo:</span>
                            <select id="sel_horas" style="margin-bottom: 0; width: 100px;">
                                <option value="1">1 Hora</option>
                                <option value="2">2 Horas</option>
                                <option value="3">3 Horas</option>
                                <option value="4">4 Horas</option>
                                <option value="6">6 Horas</option>
                                <option value="8">8 Horas</option>
                                <option value="12">12 Horas</option>
                                <option value="24">24 Horas</option>
                            </select>
                        </div>
                    </div>

                    <button type="button" class="btn btn-success" style="margin-top: 15px;" onclick="iniciarExtraccion()">🚀 INICIAR EXTRACCIÓN</button>
                </div>
            </div>
        </div>

        <div id="tab-alertas" class="tab-content" style="display: none;">
            <div class="panel" style="margin-bottom: 0; flex-shrink: 0;">
                <div style="display: flex; gap: 20px; align-items: center;">
                    <div style="flex: 1;"><label style="color: var(--warning);">ESTACIÓN A CONFIGURAR</label><select id="cfg_fuente" onchange="cargarConfigJsonEspecifico()" style="margin: 0; border-color: var(--warning);"></select></div>
                    <div style="flex: 2; color: #8b949e; font-size: 13px;">⚠️ Se guardará directamente en el archivo JSON asignado a esta fuente.</div>
                </div>
            </div>
            
            <div class="grid-2col" id="cfg_workspace" style="display: none;">
                <div class="panel">
                    <div class="panel-title">
                        <span>Horarios y Calendario</span>
                        <div class="switch-container" style="display: flex; align-items: center; gap: 10px;">
                            <span style="font-size: 12px; font-weight: bold;">ACTIVAR ALERTAS</span>
                            <label class="switch"><input type="checkbox" id="cfg_master"><span class="slider"></span></label>
                        </div>
                    </div>
                    <div class="input-group">
                        <div><label>No enviar antes de</label><input type="time" id="cfg_h_ini"></div>
                        <div><label>No enviar después de</label><input type="time" id="cfg_h_fin"></div>
                    </div>
                    <div style="display: flex; justify-content: space-between; align-items: center; margin-top: 10px; margin-bottom: 10px;">
                        <button type="button" class="btn btn-outline" onclick="cambiarMes(-1)">&#10094;</button><h3 id="mes-display" style="margin: 0; font-size: 16px;">Mes</h3><button type="button" class="btn btn-outline" onclick="cambiarMes(1)">&#10095;</button>
                    </div>
                    <div class="calendar-grid"><div class="day-name">LU</div><div class="day-name">MA</div><div class="day-name">MI</div><div class="day-name">JU</div><div class="day-name">VI</div><div class="day-name">SA</div><div class="day-name">DO</div></div>
                    <div class="calendar-grid" id="calendar-days"></div>
                </div>
                <div class="panel">
                    <div class="panel-title">Directorio de Contactos JSON</div>
                    <div class="search-box"><span class="search-icon">🔍</span><input type="text" id="buscador" placeholder="Buscar por nombre, email o grupo..." onkeyup="renderContactos()"></div>
                    <div style="background: var(--bg); padding: 12px; border-radius: 8px; margin-bottom: 15px; border: 1px solid var(--panel-light);">
                        <div style="display: flex; gap: 8px; margin-bottom: 8px;">
                            <input type="text" id="nuevo_nom" placeholder="Nombre completo" style="margin:0;">
                            <input type="text" id="nuevo_grupo" placeholder="Grupo (Ej: Jefes)" style="margin:0; width: 150px;">
                        </div>
                        <div style="display: flex; gap: 8px;">
                            <input type="text" id="nuevo_mail" placeholder="Email" style="margin:0; flex:1;">
                            <button type="button" class="btn btn-primary" onclick="addContacto()" style="padding: 0 15px;">Añadir</button>
                        </div>
                    </div>
                    <div class="group-badges" id="group-badges"></div>
                    <div class="list-box" id="lista_contactos" style="max-height: 400px; overflow-y: auto; padding-right: 5px; margin-top: 10px;"></div>
                </div>
            </div>
            <button type="button" id="btn_guardar_json" class="btn btn-success" style="display: none;" onclick="guardarConfigJsonEspecifico()">💾 SOBRESCRIBIR JSON DE LA ESTACIÓN</button>
        </div>

        <div id="tab-visor" class="tab-content" style="display: none;">
            <div class="panel" style="margin-bottom: 0; flex-shrink: 0; padding: 10px 20px;">
                <div style="display:flex; align-items:center; gap: 15px;">
                    <label style="margin:0; white-space:nowrap; font-size: 13px;">ESTACIÓN / FUENTE A EXPLORAR</label>
                    <select id="visor_fuente" onchange="cargarRaizVisor()" style="margin:0; flex:1;"></select>
                    {% if es_admin %}
                    <button type="button" class="btn btn-warning" style="color:black; margin:0;" onclick="abrirModalHistorico()">📦 Mover a Histórico</button>
                    {% endif %}
                </div>
            </div>
            
            <div class="visor-container">
                <div class="panel col-explorer">
                    <div class="panel-title" style="margin-bottom: 10px;">Archivos<button onclick="visorSubirNivel()" class="btn btn-outline" style="padding: 4px 8px; font-size:11px;">⬆ Subir</button></div>
                    <div style="font-size:11px; color:#8b949e; margin-bottom:10px; word-break: break-all;" id="visor-ruta-actual">/</div>
                    <div id="visor-lista"></div>
                </div>
                
                <div class="panel col-preview">
                    <div class="panel-title" style="display: flex; justify-content: space-between; align-items:center; margin-bottom: 10px; flex-shrink:0;">
                        <span id="visor-titulo-doc" style="white-space: nowrap; overflow: hidden; text-overflow: ellipsis; max-width: 50%;">Previsualización</span>
                        <div style="display:flex; align-items:center;">
                            <div class="zoom-controls" id="docx-zoom-bar"><button class="zoom-btn" onclick="ajustarZoom(-0.1)">-</button><span id="zoom-label" style="font-size:13px; min-width:40px; text-align:center;">100%</span><button class="zoom-btn" onclick="ajustarZoom(0.1)">+</button></div>
                            <a id="btn-descargar-visor" class="btn btn-primary" style="display:none;" href="" download>⬇ Descargar</a>
                        </div>
                    </div>
                    <div id="visor-contenido">
                        <div style="height:100%; display:flex; align-items:center; justify-content:center; color:#8b949e; width: 100%;">Selecciona un archivo a la izquierda.</div>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <div id="modal_fuentes" class="modal-overlay">
        <div class="modal">
            <div class="panel-title">Fuentes / Estaciones <button type="button" class="btn" style="background:transparent; color:white; padding:0;" onclick="cerrarModalFuentes()">✖</button></div>
            <div style="background: var(--bg); padding: 15px; border-radius: 8px; margin-bottom: 20px; border: 1px solid var(--panel-light);">
                <label>Nombre</label><input type="text" id="m_nombre" placeholder="Ej. UMA">
                <label>Ruta Resultados</label><div style="display:flex; gap:10px;"><input type="text" id="m_res" style="flex:1;"><button type="button" class="btn btn-outline" style="height: 42px;" onclick="abrirWebBrowser('m_res', 'folder')">🔍 Buscar</button></div>
                <label>Ruta Binarios</label><div style="display:flex; gap:10px;"><input type="text" id="m_bin" style="flex:1;"><button type="button" class="btn btn-outline" style="height: 42px;" onclick="abrirWebBrowser('m_bin', 'folder')">🔍 Buscar</button></div>
                <label style="color: var(--warning);">Ruta JSON</label><div style="display:flex; gap:10px;"><input type="text" id="m_json" style="border-color: var(--warning); flex:1;"><button type="button" class="btn btn-outline" style="height: 42px; border-color:var(--warning); color:var(--warning);" onclick="abrirWebBrowser('m_json', 'file_json')">🔍 Buscar</button></div>
                <hr style="border:0; border-top:1px solid #30363d; margin:15px 0;">
                <label style="color: #10b981;">CREDENCIALES DE RED (OPCIONAL)</label>
                <div style="display:flex; gap:10px; margin-bottom:0;">
                    <div style="flex:1;"><label>Usuario</label><input type="text" id="m_usr" placeholder="Ej: DSIC\\usuario" style="margin-bottom:0;"></div>
                    <div style="flex:1;"><label>Contraseña</label><input type="password" id="m_pwd" placeholder="***" style="margin-bottom:0;"></div>
                </div>
                <button type="button" class="btn btn-primary" style="width: 100%; margin-top:20px;" onclick="guardarNuevaFuente()">Guardar / Actualizar Fuente</button>
            </div>
            <div class="list-box" id="lista_fuentes_modal"></div>
        </div>
    </div>

    <div id="modal_usuarios" class="modal-overlay">
        <div class="modal">
            <div class="panel-title">Gestión de Usuarios <button type="button" class="btn" style="background:transparent; color:white; padding:0;" onclick="cerrarModalUsuarios()">✖</button></div>
            <div style="background: var(--bg); padding: 15px; border-radius: 8px; margin-bottom: 20px; border: 1px solid var(--panel-light);">
                <label>Añadir o Editar Usuario</label>
                <div style="display:flex; gap:10px; margin-bottom:10px;">
                    <input type="text" id="u_nombre" placeholder="Nombre Usuario" style="flex:1; margin:0;">
                    <input type="text" id="u_pass" placeholder="Contraseña" style="flex:1; margin:0;">
                    <select id="u_rol" style="flex:0.7; margin:0;"><option value="user">Técnico (User)</option><option value="admin">Administrador</option></select>
                </div>
                <button type="button" class="btn btn-primary" style="width: 100%;" onclick="guardarNuevoUsuario()">Añadir / Actualizar Usuario</button>
            </div>
            <div class="list-box" id="lista_usuarios_modal"></div>
        </div>
    </div>

    <div id="modal_historico" class="modal-overlay">
        <div class="modal">
            <div class="panel-title">Mover Carpeta a Histórico <button type="button" class="btn" style="background:transparent; color:white; padding:0;" onclick="cerrarModalHistorico()">✖</button></div>
            <div style="background: var(--bg); padding: 15px; border-radius: 8px; margin-bottom: 20px; border: 1px solid var(--panel-light);">
                <p style="font-size:13px; color:#8b949e;">Selecciona la carpeta que deseas archivar. Se moverá a la subcarpeta <strong>/historico</strong>.</p>
                <label>Carpeta a mover</label><select id="h_carpeta"></select>
                <button type="button" class="btn btn-warning" style="width: 100%; margin-top:15px; color:black;" onclick="moverAHistorico()">📦 Confirmar Movimiento</button>
            </div>
        </div>
    </div>

    <div id="modal_browser" class="modal-overlay">
        <div class="modal" style="width: 600px; max-width: 90%;">
            <div class="panel-title" id="mb_title">Seleccionar Ruta de Servidor <button type="button" class="btn" style="background:transparent; color:white; padding:0;" onclick="document.getElementById('modal_browser').style.display='none'">✖</button></div>
            <div style="background: var(--bg); padding: 10px; border-radius: 8px; margin-bottom: 10px; border: 1px solid var(--panel-light); display:flex; gap:10px;">
                <input type="text" id="mb_path" style="margin:0; flex:1;" placeholder="C:\\">
                <button type="button" class="btn btn-outline" onclick="mbLoadPath()">Ir</button>
                <button type="button" class="btn btn-outline" onclick="mbUp()">⬆ Subir</button>
            </div>
            <div class="list-box" id="mb_list" style="max-height: 400px; overflow-y: auto; background: var(--bg); border: 1px solid var(--panel-light);"></div>
            <div style="display:flex; justify-content: flex-end; margin-top: 15px; gap: 10px;">
                <button type="button" class="btn btn-warning" style="margin:0;" onclick="document.getElementById('modal_browser').style.display='none'">Cancelar</button>
                <button type="button" class="btn btn-success" id="mb_btn_sel" style="margin:0;" onclick="mbConfirm()">✅ Seleccionar Carpeta Actual</button>
            </div>
        </div>
    </div>

    <script>
        document.addEventListener('contextmenu', event => event.preventDefault());
        document.addEventListener('keydown', function(e) { if (e.key === 'F12' || (e.ctrlKey && e.shiftKey && e.key === 'I')) { e.preventDefault(); } });

        let fuentesGlobales = []; let usuariosGlobales = {}; let rutaVisorActual = ""; 
        let currentVisorFiles = []; let currentImagesList = []; let currentImageIndex = -1; let currentZoom = 1.0;

        window.onload = () => {
            document.getElementById('ext_f_ini').valueAsDate = new Date();
            document.getElementById('ext_f_fin').valueAsDate = new Date();
            cargarFuentesServidor();
        };

        function cambiarPestana(id, btn) {
            let contenidos = document.getElementsByClassName('tab-content');
            for(let i=0; i<contenidos.length; i++) contenidos[i].style.display = 'none';
            let botones = document.getElementsByClassName('tab-btn');
            for(let i=0; i<botones.length; i++) botones[i].classList.remove('active');
            document.getElementById(id).style.display = 'flex';
            if(btn) btn.classList.add('active');
        }

        function cargarFuentesServidor() { fetch('/api/fuentes').then(r => r.json()).then(data => { fuentesGlobales = data; renderSelectoresFuentes(); renderListaModal(); }); }
        
        function renderSelectoresFuentes() {
            const s1 = document.getElementById('ext_fuente'); const s2 = document.getElementById('cfg_fuente'); const s3 = document.getElementById('visor_fuente');
            s1.innerHTML = ''; s2.innerHTML = '<option value="">-- Selecciona Estación --</option>'; s3.innerHTML = '<option value="">-- Selecciona Estación --</option>';
            fuentesGlobales.forEach((f, i) => { 
                s1.innerHTML += `<option value="${i}">${f.nombre}</option>`; 
                s2.innerHTML += `<option value="${i}">${f.nombre}</option>`; 
                s3.innerHTML += `<option value="${i}">${f.nombre}</option>`; 
            });
            actualizarPlaceholderSalida();
        }

        let mbTargetId = ""; let mbMode = "folder";
        function abrirWebBrowser(targetId, mode) {
            mbTargetId = targetId; mbMode = mode;
            let initialPath = document.getElementById(targetId).value.trim();
            if(!initialPath) {
                const s1 = document.getElementById('ext_fuente');
                if(s1 && s1.value !== "" && fuentesGlobales[s1.value]) initialPath = fuentesGlobales[s1.value].ruta_res || "C:\\\\";
                else initialPath = "C:\\\\";
            }
            document.getElementById('mb_title').innerText = mode === 'folder' ? 'Seleccionar Carpeta (Servidor)' : 'Seleccionar JSON (Servidor)';
            document.getElementById('mb_btn_sel').style.display = mode === 'folder' ? 'block' : 'none';
            document.getElementById('modal_browser').style.display = 'flex';
            document.getElementById('mb_path').value = initialPath;
            mbLoadPath();
        }
        function mbLoadPath() {
            const p = document.getElementById('mb_path').value;
            fetch('/api/files/list', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({path: p}) }).then(r => r.json()).then(data => {
                const box = document.getElementById('mb_list'); box.innerHTML = '';
                if(data.error) { box.innerHTML = `<div style="padding:15px; color:var(--danger);">${data.error}</div>`; return; }
                data.dirs.forEach(d => { box.innerHTML += `<div class="file-item folder" onclick="mbGo('${d.full_path.replace(/\\\\/g, '\\\\\\\\')}')">📁 ${d.name}</div>`; });
                if(mbMode === 'file_json') {
                    data.files.forEach(f => {
                        if(f.name.toLowerCase().endsWith('.json')) box.innerHTML += `<div class="file-item" style="color:#10b981;" onclick="mbSelectFile('${f.full_path.replace(/\\\\/g, '\\\\\\\\')}')">📄 ${f.name}</div>`;
                        else box.innerHTML += `<div class="file-item" style="color:#8b949e;">📄 ${f.name}</div>`;
                    });
                }
            }).catch(e => { document.getElementById('mb_list').innerHTML = `<div style="padding:15px; color:var(--danger);">Error de red.</div>`; });
        }
        function mbGo(newPath) { document.getElementById('mb_path').value = newPath; mbLoadPath(); }
        function mbUp() { let p = document.getElementById('mb_path').value; let parts = p.split(/[\\\\/]/).filter(x => x); parts.pop(); let newPath = parts.join('\\\\'); if(parts.length===1 && p.includes(':\\\\')) newPath+="\\\\"; document.getElementById('mb_path').value = newPath || "C:\\\\"; mbLoadPath(); }
        function mbConfirm() { if(mbMode==='folder') { document.getElementById(mbTargetId).value = document.getElementById('mb_path').value; document.getElementById('modal_browser').style.display = 'none'; } }
        function mbSelectFile(fullPath) { document.getElementById(mbTargetId).value = fullPath; document.getElementById('modal_browser').style.display = 'none'; }

        function abrirModalFuentes() { document.getElementById('modal_fuentes').style.display = 'flex'; }
        function cerrarModalFuentes() { document.getElementById('modal_fuentes').style.display = 'none'; }
        
        function guardarNuevaFuente() {
            const f = { nombre: document.getElementById('m_nombre').value.trim(), ruta_res: document.getElementById('m_res').value.trim(), ruta_bin: document.getElementById('m_bin').value.trim(), ruta_json: document.getElementById('m_json').value.trim(), usr_red: document.getElementById('m_usr').value.trim(), pwd_red: document.getElementById('m_pwd').value.trim() };
            if(!f.nombre || !f.ruta_res) return;
            fuentesGlobales = fuentesGlobales.filter(x => x.nombre !== f.nombre); fuentesGlobales.push(f);
            fetch('/api/fuentes', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify(fuentesGlobales) }).then(() => { renderSelectoresFuentes(); renderListaModal(); document.querySelectorAll('#modal_fuentes input').forEach(i => i.value = ''); cerrarModalFuentes(); });
        }
        function eliminarFuente(nombre) { fuentesGlobales = fuentesGlobales.filter(x => x.nombre !== nombre); fetch('/api/fuentes', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify(fuentesGlobales) }).then(() => { renderSelectoresFuentes(); renderListaModal(); }); }
        function editarFuente(nombre) { const f = fuentesGlobales.find(x => x.nombre === nombre); if(f) { document.getElementById('m_nombre').value = f.nombre; document.getElementById('m_res').value = f.ruta_res; document.getElementById('m_bin').value = f.ruta_bin; document.getElementById('m_json').value = f.ruta_json; document.getElementById('m_usr').value = f.usr_red || ''; document.getElementById('m_pwd').value = f.pwd_red || ''; } }
        function renderListaModal() { const box = document.getElementById('lista_fuentes_modal'); box.innerHTML = ''; fuentesGlobales.forEach(f => { box.innerHTML += `<div class="fuente-item"><div><strong>${f.nombre}</strong><br><span style="font-size:11px; color:#8b949e;">RES: ${f.ruta_res}</span></div><div><button type="button" class="btn btn-warning" style="padding: 5px 10px; margin-right: 5px; color: black;" onclick="editarFuente('${f.nombre}')">✏️</button><button type="button" class="btn btn-danger" style="padding: 5px 10px;" onclick="eliminarFuente('${f.nombre}')">✖</button></div></div>`; }); }
        function actualizarPlaceholderSalida() { const s1 = document.getElementById('ext_fuente'); if(s1.value !== "") { document.getElementById('ext_salida').placeholder = "Auto: " + fuentesGlobales[s1.value].ruta_res + "\\\\ejecucion_a_peticion_FECHA"; } }

        function abrirModalUsuarios() { document.getElementById('modal_usuarios').style.display = 'flex'; fetch('/api/usuarios').then(r => r.json()).then(data => { if(data.error) { alert(data.error); return; } usuariosGlobales = data; renderListaUsuarios(); }); }
        function cerrarModalUsuarios() { document.getElementById('modal_usuarios').style.display = 'none'; }
        function renderListaUsuarios() { const box = document.getElementById('lista_usuarios_modal'); box.innerHTML = ''; for (const [user, info] of Object.entries(usuariosGlobales)) { const isAd = info.role === 'admin'; box.innerHTML += `<div class="fuente-item" style="border-left-color: ${isAd ? 'var(--warning)' : 'var(--accent)'}"><div><strong>${user}</strong> ${isAd ? '<span style="font-size:10px; background:var(--warning); color:black; padding:2px 5px; border-radius:4px; margin-left:5px;">ADMIN</span>' : ''}<br><span style="font-size:11px; color:#8b949e;">Pass: ${info.password}</span></div><div><button type="button" class="btn btn-warning" style="padding: 5px 10px; margin-right: 5px; color: black;" onclick="editarUsuario('${user}')">✏️</button><button type="button" class="btn btn-danger" style="padding: 5px 10px;" onclick="eliminarUsuario('${user}')">✖</button></div></div>`; } }
        function guardarNuevoUsuario() { const u = document.getElementById('u_nombre').value.trim(); const p = document.getElementById('u_pass').value.trim(); const r = document.getElementById('u_rol').value; if(!u || !p) return; usuariosGlobales[u] = {password: p, role: r}; fetch('/api/usuarios', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify(usuariosGlobales) }).then(() => { renderListaUsuarios(); document.getElementById('u_nombre').value = ''; document.getElementById('u_pass').value = ''; }); }
        function eliminarUsuario(user) { if(user === 'admin') { alert("No puedes borrar al admin principal."); return; } delete usuariosGlobales[user]; fetch('/api/usuarios', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify(usuariosGlobales) }).then(() => { renderListaUsuarios(); }); }
        function editarUsuario(user) { document.getElementById('u_nombre').value = user; document.getElementById('u_pass').value = usuariosGlobales[user].password; document.getElementById('u_rol').value = usuariosGlobales[user].role; }

        function abrirModalHistorico() {
            const idx = document.getElementById('visor_fuente').value;
            if(idx === "") { alert("Selecciona una fuente primero."); return; }
            document.getElementById('modal_historico').style.display = 'flex';
            document.getElementById('h_carpeta').innerHTML = '<option>Cargando...</option>';
            fetch('/api/files/get_folders', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({fuente: fuentesGlobales[idx]}) })
            .then(r => r.json()).then(data => {
                const sel = document.getElementById('h_carpeta'); sel.innerHTML = '';
                if(data.error) { sel.innerHTML = `<option value="">Error: ${data.error}</option>`; return; }
                if(data.dirs.length === 0) { sel.innerHTML = `<option value="">No hay carpetas</option>`; return; }
                data.dirs.forEach(d => { sel.innerHTML += `<option value="${d}">${d}</option>`; });
            }).catch(e => { document.getElementById('h_carpeta').innerHTML = `<option value="">Error de red</option>`; });
        }
        function cerrarModalHistorico() { document.getElementById('modal_historico').style.display = 'none'; }
        function moverAHistorico() {
            const idx = document.getElementById('visor_fuente').value; const carpeta = document.getElementById('h_carpeta').value;
            if(!carpeta) return; if(!confirm(`¿Seguro que quieres mover "${carpeta}" al histórico?`)) return;
            mostrarLoading("Moviendo carpeta...");
            fetch('/api/files/move_history', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({fuente: fuentesGlobales[idx], folder_name: carpeta}) })
            .then(r => r.json()).then(data => { ocultarLoading(); if(data.error) alert("Error: " + data.error); else { alert("Carpeta archivada."); cerrarModalHistorico(); cargarRaizVisor(); } }).catch(e => { ocultarLoading(); alert("Error de red."); });
        }

        function cargarRaizVisor() { const idx = document.getElementById('visor_fuente').value; if(idx === "") { document.getElementById('visor-lista').innerHTML = ''; return; } rutaVisorActual = fuentesGlobales[idx].ruta_res; listarDirectorio(rutaVisorActual); }
        function visorSubirNivel() { const idx = document.getElementById('visor_fuente').value; if(idx === "") return; const raiz = fuentesGlobales[idx].ruta_res; if(rutaVisorActual.length <= raiz.length) return; let parts = rutaVisorActual.split(/[\\\\/]/); parts.pop(); rutaVisorActual = parts.join('\\\\'); listarDirectorio(rutaVisorActual); }
        
        function listarDirectorio(path) {
            document.getElementById('visor-ruta-actual').innerText = path; currentVisorFiles = []; currentImagesList = []; currentImageIndex = -1;
            const idx = document.getElementById('visor_fuente').value; const f_data = fuentesGlobales[idx];
            fetch('/api/files/list', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({path: path, fuente: f_data}) })
            .then(r => r.json()).then(data => {
                const box = document.getElementById('visor-lista'); box.innerHTML = '';
                if(data.error) { box.innerHTML = `<span style="color:var(--danger);">${data.error}</span>`; return; }
                data.dirs.forEach(d => { box.innerHTML += `<div class="file-item folder" onclick="navegarACarpeta('${d.full_path.replace(/\\\\/g, '\\\\\\\\')}')">📁 ${d.name}</div>`; });
                data.files.forEach(f => {
                    currentVisorFiles.push(f); let icon = "📄"; let ext = f.name.split('.').pop().toLowerCase();
                    if(['png','jpg','jpeg'].includes(ext)) { icon = "🖼️"; currentImagesList.push(f); } if(['csv','xlsx'].includes(ext)) icon = "📊"; if(ext === 'pdf') icon = "📕"; if(ext === 'docx') icon = "📘";
                    box.innerHTML += `<div class="file-item" id="file-${f.name.replace(/[^a-zA-Z0-9]/g,'')}" onclick="verArchivo('${f.full_path.replace(/\\\\/g, '\\\\\\\\')}', '${f.name}')">${icon} ${f.name}</div>`;
                });
            }).catch(e => { document.getElementById('visor-lista').innerHTML = `<span style="color:var(--danger);">Error de conexión.</span>`; });
        }
        function navegarACarpeta(newPath) { rutaVisorActual = newPath; listarDirectorio(rutaVisorActual); }

        document.addEventListener('keydown', (e) => { if(document.getElementById('tab-visor').style.display !== 'none' && currentImagesList.length > 0 && currentImageIndex !== -1) { if(e.key === 'ArrowLeft' && currentImageIndex > 0) renderImagePreview(currentImageIndex - 1); if(e.key === 'ArrowRight' && currentImageIndex < currentImagesList.length - 1) renderImagePreview(currentImageIndex + 1); } });

        function renderImagePreview(index) {
            currentImageIndex = index; let imgData = currentImagesList[index]; document.getElementById('visor-titulo-doc').innerText = imgData.name; document.getElementById('btn-descargar-visor').style.display = 'block'; document.getElementById('btn-descargar-visor').href = `/api/files/download?path=${encodeURIComponent(imgData.full_path)}`; document.getElementById('docx-zoom-bar').style.display = 'none';
            document.querySelectorAll('.file-item').forEach(el => el.classList.remove('active')); let activeEl = document.getElementById(`file-${imgData.name.replace(/[^a-zA-Z0-9]/g,'')}`); if(activeEl) activeEl.classList.add('active');
            let html = `<div style="position:relative; width:100%; height:100%; display:flex; justify-content:center; align-items:center;">`; if(currentImageIndex > 0) html += `<button onclick="renderImagePreview(${currentImageIndex - 1})" class="nav-img-btn nav-left">❮</button>`; html += `<img src="/api/files/serve?path=${encodeURIComponent(imgData.full_path)}" style="max-width:100%; max-height:100%; object-fit:contain; border-radius:4px;">`; if(currentImageIndex < currentImagesList.length - 1) html += `<button onclick="renderImagePreview(${currentImageIndex + 1})" class="nav-img-btn nav-right">❯</button>`; html += `</div>`;
            let cont = document.getElementById('visor-contenido'); cont.style.padding = "0"; cont.style.alignItems = "center"; cont.innerHTML = html;
        }
        function ajustarZoom(delta) { currentZoom += delta; if(currentZoom < 0.3) currentZoom = 0.3; if(currentZoom > 3.0) currentZoom = 3.0; document.getElementById('zoom-label').innerText = Math.round(currentZoom * 100) + '%'; let docxContainer = document.getElementById('docx-container'); if(docxContainer) { docxContainer.style.transform = `scale(${currentZoom})`; } }

        function verArchivo(fullPath, fileName) {
            document.querySelectorAll('.file-item').forEach(el => el.classList.remove('active')); let activeEl = document.getElementById(`file-${fileName.replace(/[^a-zA-Z0-9]/g,'')}`); if(activeEl) activeEl.classList.add('active');
            let ext = fileName.split('.').pop().toLowerCase(); document.getElementById('docx-zoom-bar').style.display = 'none';
            if(['png','jpg','jpeg'].includes(ext)) { let idx = currentImagesList.findIndex(f => f.name === fileName); if(idx !== -1) renderImagePreview(idx); return; }
            currentImageIndex = -1; document.getElementById('visor-titulo-doc').innerText = fileName; const btnDescarga = document.getElementById('btn-descargar-visor'); btnDescarga.style.display = 'block'; btnDescarga.href = `/api/files/download?path=${encodeURIComponent(fullPath)}`;
            const cont = document.getElementById('visor-contenido'); cont.innerHTML = '<div style="height:100%; display:flex; align-items:center; justify-content:center;"><div class="loader"></div></div>';
            
            if (ext === 'pdf') { cont.style.padding = "0"; cont.innerHTML = `<iframe src="/api/files/serve?path=${encodeURIComponent(fullPath)}" style="width:100%; height:100%; border:none;"></iframe>`; }
            else if (['csv', 'xlsx'].includes(ext)) { fetch('/api/files/preview', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({path: fullPath}) }).then(r => r.json()).then(data => { cont.style.padding = "10px"; cont.style.alignItems = "flex-start"; cont.innerHTML = `<div style="width:100%; background:white; padding:15px; border-radius:8px; overflow:hidden;">${data.html}</div>`; $('#tabla-preview').DataTable({ scrollX: true, pageLength: 25, language: { url: '//cdn.datatables.net/plug-ins/1.13.6/i18n/es-ES.json' } }); }).catch(e => { cont.innerHTML = `<span style="color:var(--danger);">Error al cargar vista previa.</span>`; }); }
            else if (ext === 'docx') { fetch('/api/files/preview', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({path: fullPath}) }).then(r => r.json()).then(data => { cont.style.padding = "20px"; cont.style.alignItems = "flex-start"; currentZoom = 1.0; document.getElementById('zoom-label').innerText = '100%'; document.getElementById('docx-zoom-bar').style.display = 'flex'; cont.innerHTML = `<div id="docx-container" style="width:100%; background:white; padding:40px; border-radius:4px; color:black; transform-origin: top center; transition: transform 0.2s ease-out;">${data.html}</div>`; }).catch(e => { cont.innerHTML = `<span style="color:var(--danger);">Error al cargar vista previa.</span>`; }); }
            else { cont.style.padding = "20px"; cont.style.alignItems = "center"; cont.innerHTML = `<div style="height:100%; display:flex; align-items:center; justify-content:center;"><p>Vista previa no disponible para este formato. Descárgalo para verlo.</p></div>`; }
        }

        // --- EXTRACCIÓN ---
        function iniciarExtraccion() {
            const idx = document.getElementById('ext_fuente').value; if(idx === "") { alert("Selecciona una fuente."); return; }
            const payload = { fuente: fuentesGlobales[idx], origen_datos: document.querySelector('input[name="origen_datos"]:checked').value, f_ini: document.getElementById('ext_f_ini').value, h_ini: document.getElementById('ext_h_ini').value, f_fin: document.getElementById('ext_f_fin').value, h_fin: document.getElementById('ext_h_fin').value, salida_manual: document.getElementById('ext_salida').value, entregables: { graf_2dt: document.getElementById('chk_2dt').checked, graf_2df: document.getElementById('chk_2df').checked, graf_3d: document.getElementById('chk_3d').checked, csv: document.getElementById('chk_csv').checked, excel: document.getElementById('chk_excel').checked, word: document.getElementById('chk_word').checked, dividir: document.getElementById('chk_dividir').checked, horas_div: parseInt(document.getElementById('sel_horas').value) }, enviar_email: document.getElementById('chk_email').checked, emails: document.getElementById('ext_emails').value };
            mostrarLoading("Ejecutando orquestador matemático... Mira la consola CMD para detalles.");
            fetch('/api/extraer', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify(payload) })
            .then(r => r.json()).then(resp => { 
                ocultarLoading(); 
                if(resp.status === "ok") { 
                    alert("✅ EXTRACCIÓN COMPLETADA.\\nGuardado en:\\n" + resp.ruta_salida); 
                    if(payload.enviar_email) alert("Emails enviados al servidor de correo."); 
                } else {
                    alert("❌ ERROR:\\n" + resp.error); 
                }
            }).catch(e => { 
                ocultarLoading(); alert("Error de red con el servidor."); 
            });
        }
        function mostrarLoading(txt) { document.getElementById('loading-text').innerHTML = txt; document.getElementById('overlay-loading').style.display = 'flex'; }
        function ocultarLoading() { document.getElementById('overlay-loading').style.display = 'none'; }
        
        // --- CONTACTOS DE ALERTA ---
        function getColorForGroup(g) { const colors = ['#3b82f6', '#ef4444', '#10b981', '#f59e0b', '#8b5cf6', '#ec4899', '#14b8a6', '#f97316']; let hash = 0; for (let i = 0; i < g.length; i++) hash = g.charCodeAt(i) + ((hash << 5) - hash); return colors[Math.abs(hash) % colors.length]; }
        function addContacto() { const n = document.getElementById('nuevo_nom').value; const e = document.getElementById('nuevo_mail').value.toLowerCase(); let g = document.getElementById('nuevo_grupo').value.trim(); if(!g) g = "General"; if(e && !configJsonActivo.contactos.some(c => c.email === e)) { configJsonActivo.contactos.push({nombre: n||e.split('@')[0], email: e, grupo: g, activo: true}); document.getElementById('nuevo_nom').value = ''; document.getElementById('nuevo_mail').value = ''; document.getElementById('nuevo_grupo').value = ''; renderContactos(); } }
        function toggleContacto(email) { let c = configJsonActivo.contactos.find(x => x.email === email); if(c) { c.activo = !c.activo; renderContactos(); } }
        function delContacto(email) { configJsonActivo.contactos = configJsonActivo.contactos.filter(x => x.email !== email); renderContactos(); }
        window.toggleGrupoMasivo = function(grupo) { let estado = true; if(grupo === 'ALL') { estado = configJsonActivo.contactos.some(c => !c.activo); configJsonActivo.contactos.forEach(c => c.activo = estado); } else { let miembros = configJsonActivo.contactos.filter(c => c.grupo === grupo); estado = miembros.some(c => !c.activo); configJsonActivo.contactos.forEach(c => { if(c.grupo === grupo) c.activo = estado; }); } renderContactos(); };
        
        function editarModificarContacto(email) {
            let c = configJsonActivo.contactos.find(x => x.email === email);
            if(c) {
                document.getElementById('nuevo_nom').value = c.nombre;
                document.getElementById('nuevo_mail').value = c.email;
                document.getElementById('nuevo_grupo').value = c.grupo;
                delContacto(email);
            }
        }
        function renderContactos() { const box = document.getElementById('lista_contactos'); const search = document.getElementById('buscador').value.toLowerCase(); const badgesBox = document.getElementById('group-badges'); box.innerHTML = ''; badgesBox.innerHTML = ''; const gruposUnicos = [...new Set(configJsonActivo.contactos.map(c => c.grupo))]; if (gruposUnicos.length > 0) { badgesBox.innerHTML = `<button type="button" class="badge-btn" onclick="window.toggleGrupoMasivo('ALL')" style="border-color:var(--accent);"><span style="width:8px; height:8px; border-radius:50%; background:var(--accent);"></span> TODOS</button>`; gruposUnicos.forEach(g => { const cColor = getColorForGroup(g); badgesBox.innerHTML += `<button type="button" class="badge-btn" onclick="window.toggleGrupoMasivo('${g}')" style="border-color:${cColor};"><span style="width:8px; height:8px; border-radius:50%; background:${cColor};"></span> ${g}</button>`; }); } let filtrados = configJsonActivo.contactos.filter(c => c.nombre.toLowerCase().includes(search) || c.email.toLowerCase().includes(search) || c.grupo.toLowerCase().includes(search)); if (filtrados.length === 0) { box.innerHTML = '<div style="text-align:center; padding:20px; color:#8b949e; font-style:italic;">No hay resultados.</div>'; return; } filtrados.forEach((c) => { const cColor = getColorForGroup(c.grupo); box.innerHTML += `<div class="contact-card ${c.activo ? '' : 'disabled'}"><div style="flex:1;"><div style="display:flex; align-items:center; gap: 10px;"><input type="checkbox" style="width:16px; height:16px;" ${c.activo?'checked':''} onchange="toggleContacto('${c.email}')"><span style="font-weight:bold; color:white; font-size:14px;">${c.nombre}</span><span class="contact-group" style="color:${cColor}; border-color:${cColor}; margin-left:5px;">${c.grupo}</span></div><div style="font-size:12px; color:#c9d1d9; margin-top:5px; margin-left: 26px;">${c.email}</div></div><div style="display:flex; gap:5px;"><button type="button" class="btn btn-warning" style="padding: 5px 10px; color: black;" onclick="editarModificarContacto('${c.email}')">✏️</button><button type="button" class="btn btn-danger" style="padding: 5px 10px;" onclick="delContacto('${c.email}')">✖</button></div></div>`; }); }

        // --- CALENDARIO Y CONFIG ---
        let mesMostrado = new Date().getMonth();
        let anoMostrado = new Date().getFullYear();
        let fechasBloqueadas = [];
        
        function fStr(a,m,d) { return `${a}-${String(m+1).padStart(2,'0')}-${String(d).padStart(2,'0')}`; }
        function renderCalendario() { const meses = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]; document.getElementById('mes-display').innerText = `${meses[mesMostrado]} ${anoMostrado}`; const grid = document.getElementById('calendar-days'); grid.innerHTML = ''; let pDia = new Date(anoMostrado, mesMostrado, 1).getDay(); pDia = pDia === 0 ? 6 : pDia - 1; let dMes = new Date(anoMostrado, mesMostrado + 1, 0).getDate(); for (let i = 0; i < pDia; i++) grid.innerHTML += `<div class="day empty"></div>`; for (let d = 1; d <= dMes; d++) { let strF = fStr(anoMostrado, mesMostrado, d); let isBlocked = fechasBloqueadas.includes(strF); grid.innerHTML += `<div class="day ${isBlocked?'blocked':'active'}" onclick="toggleDia('${strF}')">${d}</div>`; } }
        function cambiarMes(s) { mesMostrado += s; if(mesMostrado<0){mesMostrado=11;anoMostrado--;} else if(mesMostrado>11){mesMostrado=0;anoMostrado++;} renderCalendario(); }
        function toggleDia(str) { let i = fechasBloqueadas.indexOf(str); if(i===-1) fechasBloqueadas.push(str); else fechasBloqueadas.splice(i,1); renderCalendario(); }

        let configJsonActivo = {};
        function cargarConfigJsonEspecifico() { 
            const idx = document.getElementById('cfg_fuente').value; 
            if(idx === "") { document.getElementById('cfg_workspace').style.display = 'none'; document.getElementById('btn_guardar_json').style.display = 'none'; return; } 
            mostrarLoading("Leyendo JSON..."); 
            fetch('/api/config', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({ruta_json: fuentesGlobales[idx].ruta_json, fuente: fuentesGlobales[idx]}) })
            .then(r => r.json()).then(data => { 
                configJsonActivo = data; document.getElementById('cfg_master').checked = data.alertas_activas || false; document.getElementById('cfg_h_ini').value = data.hora_inicio || "00:00"; document.getElementById('cfg_h_fin').value = data.hora_fin || "23:59"; fechasBloqueadas = data.fechas_bloqueadas || []; if(!configJsonActivo.contactos) configJsonActivo.contactos = []; configJsonActivo.contactos.forEach(c => { if(!c.grupo) c.grupo = "General"; }); document.getElementById('cfg_workspace').style.display = 'grid'; document.getElementById('btn_guardar_json').style.display = 'block'; renderCalendario(); renderContactos(); ocultarLoading(); 
            }).catch(e => {
                console.error(e);
                ocultarLoading(); 
                alert("JS ERROR TRACE:\\n" + e.message + "\\n\\n" + e.stack);
            }); 
        }
        function guardarConfigJsonEspecifico() { const idx = document.getElementById('cfg_fuente').value; const fuente = fuentesGlobales[idx]; configJsonActivo.alertas_activas = document.getElementById('cfg_master').checked; configJsonActivo.hora_inicio = document.getElementById('cfg_h_ini').value; configJsonActivo.hora_fin = document.getElementById('cfg_h_fin').value; configJsonActivo.fechas_bloqueadas = fechasBloqueadas; mostrarLoading("Guardando..."); fetch('/api/config/save', { method: 'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({ruta_json: fuente.ruta_json, fuente: fuente, datos: configJsonActivo}) }).then(r => r.json()).then(resp => { ocultarLoading(); if(resp.status === "ok") alert("JSON Guardado."); else alert("Error."); }).catch(e => { ocultarLoading(); alert("Error conectando al servidor");}); }
    </script>
</body>
</html>
"""

# ==============================================================================
# RUTAS DEL BACKEND (FLASK Y CÁLCULOS)
# ==============================================================================
@app.route('/')
def home():
    es_admin = session.get('rol') == 'admin'
    return render_template_string(HTML_TEMPLATE, es_admin=es_admin)

@app.route('/favicon.ico')
def favicon():
    return send_file('analisis-de-datos.ico', mimetype='image/vnd.microsoft.icon')

@app.route('/api/usuarios', methods=['GET', 'POST'])
def api_usuarios():
    if session.get('rol') != 'admin': return jsonify({"error": "No autorizado"}), 403
    if request.method == 'GET': return jsonify(cargar_usuarios())
    else:
        with open(ARCHIVO_USUARIOS, 'w', encoding='utf-8') as f: json.dump(request.json, f, indent=4, ensure_ascii=False)
        return jsonify({"status": "ok"})

@app.route('/api/browse_folder')
def browse_folder():
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.attributes('-topmost', True)
        root.withdraw()
        folder = filedialog.askdirectory(initialdir='C:\\', title='Selecciona la carpeta')
        root.destroy()
        return jsonify({"folder": folder.replace('/', '\\') if folder else ""})
    except: return jsonify({"folder": ""})

@app.route('/api/browse_file')
def browse_file():
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.attributes('-topmost', True)
        root.withdraw()
        file_path = filedialog.askopenfilename(initialdir='C:\\', title='Selecciona el archivo JSON', filetypes=[('JSON Files', '*.json'), ('All Files', '*.*')])
        root.destroy()
        return jsonify({"file": file_path.replace('/', '\\') if file_path else ""})
    except: return jsonify({"file": ""})

@app.route('/api/files/list', methods=['POST'])
def api_files_list():
    path = request.json.get('path')
    fuente = request.json.get('fuente', {})
    if fuente: conectar_red_windows(path, fuente.get('usr_red'), fuente.get('pwd_red'))
    
    if not path or not os.path.exists(path): return jsonify({"error": "Ruta de red inaccesible o no existe."})
    dirs, files = [], []
    try:
        for item in os.listdir(path):
            full_path = os.path.join(path, item)
            if os.path.isdir(full_path): dirs.append({"name": item, "full_path": full_path})
            else: files.append({"name": item, "full_path": full_path})
    except Exception as e: return jsonify({"error": str(e)})
    return jsonify({"dirs": sorted(dirs, key=lambda x: x['name']), "files": sorted(files, key=lambda x: x['name'])})

@app.route('/api/files/get_folders', methods=['POST'])
def api_get_folders():
    if session.get('rol') != 'admin': return jsonify({"error": "No autorizado"}), 403
    try:
        fuente = request.json.get('fuente')
        ruta_res = fuente['ruta_res']
        conectar_red_windows(ruta_res, fuente.get('usr_red'), fuente.get('pwd_red'))
        if not os.path.exists(ruta_res): return jsonify({"dirs": []})
        dirs = [d for d in os.listdir(ruta_res) if os.path.isdir(os.path.join(ruta_res, d)) and d.lower() != 'historico']
        return jsonify({"dirs": sorted(dirs)})
    except Exception as e: return jsonify({"error": str(e)})

@app.route('/api/files/move_history', methods=['POST'])
def api_move_history():
    if session.get('rol') != 'admin': return jsonify({"error": "No autorizado"}), 403
    try:
        fuente = request.json.get('fuente')
        folder_name = request.json.get('folder_name')
        ruta_res = fuente['ruta_res']
        
        conectar_red_windows(ruta_res, fuente.get('usr_red'), fuente.get('pwd_red'))
        
        source = os.path.join(ruta_res, folder_name)
        hist_dir = os.path.join(ruta_res, 'historico')
        dest = os.path.join(hist_dir, folder_name)
        
        if not os.path.exists(source): return jsonify({"error": "La carpeta origen ya no existe."})
        os.makedirs(hist_dir, exist_ok=True)
        shutil.move(source, dest)
        return jsonify({"status": "ok"})
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route('/api/files/serve')
def api_files_serve(): return send_file(request.args.get('path'))

@app.route('/api/files/download')
def api_files_download(): return send_file(request.args.get('path'), as_attachment=True)

@app.route('/api/files/preview', methods=['POST'])
def api_files_preview():
    path = request.json.get('path')
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == '.csv':
            df = pd.read_csv(path, sep=';', decimal=',', on_bad_lines='skip', nrows=500, encoding='utf-8', comment='#')
            return jsonify({"html": df.to_html(table_id="tabla-preview", classes='display nowrap', index=False)})
        elif ext == '.xlsx':
            df = pd.read_excel(path, nrows=500)
            return jsonify({"html": df.to_html(table_id="tabla-preview", classes='display nowrap', index=False)})
        elif ext == '.docx':
            if HAS_MAMMOTH:
                with open(path, "rb") as docx_file: result = mammoth.convert_to_html(docx_file)
                return jsonify({"html": result.value})
            elif HAS_DOCX:
                doc = Document(path)
                return jsonify({"html": f"<b>[Aviso: Instala 'mammoth' en consola para ver formato real]</b><br><br>" + "<br>".join([p.text for p in doc.paragraphs if p.text.strip()])})
            else: return jsonify({"html": "Librería python-docx o mammoth no instalada."})
        elif ext in ['.txt', '.log', '.json']:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f: return jsonify({"html": f"<pre style='color:#a9dbcf; white-space:pre-wrap;'>{f.read(10000)}</pre>"})
        else: return jsonify({"html": "Formato no soportado para previsualización."})
    except Exception as e: return jsonify({"html": f"<span style='color:red;'>Error: {e}</span>"})

def cargar_fuentes_globales():
    if os.path.exists(ARCHIVO_FUENTES_GLOBAL):
        try:
            with open(ARCHIVO_FUENTES_GLOBAL, 'r', encoding='utf-8') as f: return json.load(f)
        except: pass
    return []

@app.route('/api/fuentes', methods=['GET', 'POST'])
def api_fuentes():
    if request.method == 'GET': return jsonify(cargar_fuentes_globales())
    else:
        with open(ARCHIVO_FUENTES_GLOBAL, 'w', encoding='utf-8') as f: json.dump(request.json, f, indent=4, ensure_ascii=False)
        return jsonify({"status": "ok"})

@app.route('/api/config', methods=['POST'])
def api_get_config():
    try:
        req = request.json or {}
        ruta = req.get('ruta_json') or ""
        
        if ruta and os.path.exists(ruta):
            try:
                with open(ruta, 'r', encoding='utf-8') as f: return jsonify(json.load(f))
            except: pass
            
        return jsonify({"alertas_activas": False, "hora_inicio": "00:00", "hora_fin": "23:59", "fechas_bloqueadas": [], "contactos": []})
    except Exception as e:
        return jsonify({"alertas_activas": False, "hora_inicio": "00:00", "hora_fin": "23:59", "fechas_bloqueadas": [], "contactos": [], "error_server": str(e)})

@app.route('/api/config/save', methods=['POST'])
def api_save_config():
    try:
        req = request.json or {}
        ruta_json = req.get('ruta_json') or ""
        
        if not ruta_json:
            return jsonify({"status": "error", "error": "Ruta de destino no configurada o vacía."})
            
        os.makedirs(os.path.dirname(os.path.abspath(ruta_json)), exist_ok=True)
        with open(ruta_json, 'w', encoding='utf-8') as f: json.dump(req.get('datos'), f, indent=4, ensure_ascii=False)
        return jsonify({"status": "ok"})
    except Exception as e: return jsonify({"status": "error", "error": str(e)})

@app.route('/api/extraer', methods=['POST'])
def api_extraer():
    try:
        payload = request.json
        fuente = payload['fuente']
        ent = payload['entregables']
        
        print("\n" + "="*60)
        print(f"🚀 INICIANDO EXTRACCIÓN - MODO: {payload['origen_datos'].upper()}")
        print("="*60)
        print(f"📍 Fuente: {fuente['nombre']}")
        
        conectar_red_windows(fuente['ruta_res'], fuente.get('usr_red'), fuente.get('pwd_red'))
        
        dt_inicio = datetime.datetime.strptime(f"{payload['f_ini']} {payload['h_ini']}", "%Y-%m-%d %H:%M")
        dt_fin = datetime.datetime.strptime(f"{payload['f_fin']} {payload['h_fin']}", "%Y-%m-%d %H:%M").replace(second=59)
        print(f"🕒 Rango Solicitado: {dt_inicio} -> {dt_fin}")
        sys.stdout.flush()
        
        titulo_graficas = f"{dt_inicio.strftime('%d/%m/%Y %H:%M')} a {dt_fin.strftime('%d/%m/%Y %H:%M')}"
        
        carpeta_auto = f"ejecucion_a_peticion_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}" # Segundos para id único
        
        ruta_salida_final = payload.get('salida_manual', '').strip()
        if not ruta_salida_final:
            base_res = os.path.abspath(fuente['ruta_res'])
            ruta_salida_final = os.path.join(base_res, carpeta_auto)
        else:
            base_res = os.path.abspath(ruta_salida_final)
            ruta_salida_final = os.path.join(base_res, carpeta_auto)

        if str(base_res).startswith('\\\\'): 
            conectar_red_windows(base_res, fuente.get('usr_red'), fuente.get('pwd_red'))
            es_red = True
        else:
            es_red = False
                
        # Para evitar cortes de red y FileNotFound con matplotlib/PIL en discos SMB, trabajamos en local primero
        if es_red:
            ruta_salida = os.path.abspath(f"temp_extract_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}")
            print(f"📂 Carpeta de trabajo LOCAL (Temporal): {ruta_salida}")
        else:
            ruta_salida = ruta_salida_final
            print(f"📂 Carpeta Destino: {ruta_salida}")
            
        try:
            os.makedirs(ruta_salida, exist_ok=True)
        except Exception as e_dir:
            print(f"⚠️ Aviso al crear carpeta base: {e_dir}")
            
        sys.stdout.flush()
        
        data = []
        
        # --- EXTRACCIÓN BINARIOS ---
        if payload['origen_datos'] == 'binario':
            print(f"🔍 Buscando binarios en: {fuente['ruta_bin']}")
            conectar_red_windows(fuente['ruta_bin'], fuente.get('usr_red'), fuente.get('pwd_red'))
            ruta_bin = fuente['ruta_bin']
            
            if not os.path.exists(ruta_bin): 
                print(f"❌ ERROR: La ruta de binarios no existe o es inaccesible.")
                return jsonify({"status": "error", "error": f"Ruta inaccesible: {ruta_bin}"})
                
            archivos_a_procesar = []
            ts_min = (dt_inicio - datetime.timedelta(days=1)).timestamp()
            ts_max = (dt_fin + datetime.timedelta(days=1)).timestamp()
            
            for f in os.listdir(ruta_bin):
                ruta_completa = os.path.join(ruta_bin, f)
                if os.path.isfile(ruta_completa) and not os.path.splitext(f)[1]:
                    mtime = os.path.getmtime(ruta_completa)
                    ctime = os.path.getctime(ruta_completa)
                    if (ts_min <= mtime <= ts_max) or (ts_min <= ctime <= ts_max):
                        archivos_a_procesar.append(ruta_completa)
                        
            print(f"🎯 Binarios en rango de fecha: {len(archivos_a_procesar)}")
            sys.stdout.flush()
            
            if not archivos_a_procesar: 
                return jsonify({"status": "error", "error": "No hay binarios con fechas cercanas a las solicitadas."})
            
            archivos_a_procesar.sort(key=lambda x: os.path.getmtime(x))
            puntos_totales = 0
            
            for ar in archivos_a_procesar:
                print(f"⚙️ Procesando binario: {os.path.basename(ar)}")
                sys.stdout.flush()
                try:
                    tamano = os.path.getsize(ar)
                    if tamano <= 26: continue
                    with open(ar, 'rb') as f:
                        head = f.read(512)
                        texto_head = ""
                        try:
                            texto_head = head.decode('utf-16-le', errors='ignore')
                        except: pass
                        if any(palabra in texto_head for palabra in ["CTER", "EA-MALAGA", "ALARMAS", "Modo de med"]):
                            continue
                        if head.startswith(b'\x00' * 32):
                            continue
                            
                        f.seek(2)
                        while True:
                            c = f.read(26 * 10000)
                            if not c: break
                            for i in range(0, len(c), 26):
                                b = c[i:i+26]
                                if len(b) < 26: break
                                try:
                                    # ORDEN ORIGINAL DE LA APLICACIÓN DE ESCRITORIO
                                    tr, fr, lv, ex = struct.unpack('<QddH', b)
                                    if abs(lv) > 300 or math.isnan(fr) or math.isnan(lv): continue 
                                    if tr < 130000000000000000 or tr > 150000000000000000: continue
                                    # Fix Zona Horaria: Conversión automatizada UTC a hora Local de España
                                    ts_utc = datetime.datetime(1601, 1, 1) + datetime.timedelta(microseconds=tr/10)
                                    ts = ts_utc.replace(tzinfo=datetime.timezone.utc).astimezone().replace(tzinfo=None)
                                    if dt_inicio <= ts <= dt_fin and 2020 < ts.year < 2100:
                                        data.append((ts, fr/1e6, lv))
                                        puntos_totales += 1
                                except Exception: pass
                except Exception as e_file: print(f"⚠️ Error {ar}: {e_file}")
            print(f"✅ Extracción binaria finalizada. Puntos útiles: {puntos_totales}")

        # --- EXTRACCIÓN CSV ---
        else:
            ruta_busqueda = fuente['ruta_res']
            print(f"🔍 Buscando CSVs recursivamente en: {ruta_busqueda}")
            sys.stdout.flush()
            
            if not os.path.exists(ruta_busqueda): return jsonify({"status": "error", "error": f"La ruta no existe: {ruta_busqueda}"})
            
            archivos_csv = []
            
            # 1. Crear permutaciones de los días solicitados (ej. 260401, 01_04_2026, 20260401)
            valid_date_strings = []
            curr_d = dt_inicio.date()
            while curr_d <= dt_fin.date():
                valid_date_strings.extend([
                    curr_d.strftime('%y%m%d'),      # 260401
                    curr_d.strftime('%Y%m%d'),      # 20260401
                    curr_d.strftime('%d_%m_%Y')     # 01_04_2026
                ])
                curr_d += datetime.timedelta(days=1)
            
            for root, dirs, files in os.walk(ruta_busqueda):
                for f in files:
                    if f.endswith('.csv') and not f.startswith('~'):
                        ruta_csv = os.path.join(root, f)
                        
                        # Filtro instantáneo de red: Si la fecha está en el nombre del fichero o carpeta, entra seguro.
                        if any(ds in ruta_csv for ds in valid_date_strings):
                            archivos_csv.append(ruta_csv)
                        else:
                            # Filtro profundo superficial: Lee rápido y localmente la 1ª o 2ª línea sin levantar pandas
                            try:
                                with open(ruta_csv, 'r', encoding='utf-8') as f_read:
                                    for row in f_read:
                                        if row.strip() and not row.startswith('#') and 'Tiempo' not in row and 'Fecha' not in row and 'T' not in row:
                                            ds = row.split(';')[0].split(',')[0].strip(' "') # Primer valor = Date
                                            try:
                                                dt_csv = pd.to_datetime(ds, dayfirst=True)
                                                if pd.notnull(dt_csv) and (dt_inicio.date() <= dt_csv.date() <= dt_fin.date()):
                                                    archivos_csv.append(ruta_csv)
                                            except: pass
                                            break 
                            except: pass
            
            print(f"📄 Total de CSV pre-filtrados para este rango: {len(archivos_csv)}")
            sys.stdout.flush()
            
            if not archivos_csv: return jsonify({"status": "error", "error": "No se encontraron CSV en esas fechas concretas."})

            data_frames = []
            for idx, csv_file in enumerate(archivos_csv):
                try:
                    df_temp = pd.read_csv(csv_file, sep=';', decimal=',', on_bad_lines='skip', encoding='utf-8', comment='#')
                    
                    col_tiempo = 'Tiempo' if 'Tiempo' in df_temp.columns else ('T' if 'T' in df_temp.columns else None)
                    col_freq = 'Frecuencia (Hz)' if 'Frecuencia (Hz)' in df_temp.columns else ('F' if 'F' in df_temp.columns else None)
                    col_level = 'Level' if 'Level' in df_temp.columns else ('L' if 'L' in df_temp.columns else None)

                    if col_tiempo and col_freq and col_level:
                        # Fix de comas a puntos y parseo
                        tiempos_str = df_temp[col_tiempo].astype(str).str.replace(',', '.')
                        # Intentamos el formato original del bot (día/mes/año)
                        df_temp[col_tiempo] = pd.to_datetime(tiempos_str, format='%d/%m/%Y %H:%M:%S.%f', errors='coerce')
                        
                        # Si fallan algunos, intentamos el parseo automatico (fallback)
                        if df_temp[col_tiempo].isna().all():
                            df_temp[col_tiempo] = pd.to_datetime(tiempos_str, dayfirst=True, errors='coerce')

                        mask = (df_temp[col_tiempo] >= dt_inicio) & (df_temp[col_tiempo] <= dt_fin)
                        df_filtrado = df_temp.loc[mask].copy()
                        
                        if not df_filtrado.empty:
                            df_filtrado[col_level] = pd.to_numeric(df_filtrado[col_level].astype(str).str.replace(',', '.'), errors='coerce')
                            df_filtrado[col_freq] = pd.to_numeric(df_filtrado[col_freq].astype(str).str.replace(',', '.'), errors='coerce')
                            
                            df_filtrado = df_filtrado.rename(columns={col_tiempo: 'T', col_freq: 'F', col_level: 'L'})
                            if df_filtrado['F'].mean() > 10000: df_filtrado['F'] = df_filtrado['F'] / 1e6
                            data_frames.append(df_filtrado[['T', 'F', 'L']])
                            
                            print(f"✔️ [{idx+1}/{len(archivos_csv)}] {os.path.basename(csv_file)}: Aportó {len(df_filtrado)} filas útiles.")
                            sys.stdout.flush()
                except Exception as e: pass
            
            if not data_frames: 
                print("❌ Ningún CSV contenía datos para esa franja horaria.")
                sys.stdout.flush()
                return jsonify({"status": "error", "error": "Ningún CSV contiene datos para esa fecha exacta."})
            
            df_concat = pd.concat(data_frames, ignore_index=True)
            df_concat = df_concat.dropna().sort_values('T')
            data = list(df_concat.itertuples(index=False, name=None))
            print(f"✅ CSV unificados. Puntos útiles totales: {len(data)}")
            sys.stdout.flush()

        if not data: 
            return jsonify({"status": "error", "error": "Los datos recuperados están vacíos tras aplicar los filtros."})

        # --- DIBUJADO Y EXPORTACIÓN ---
        df = pd.DataFrame(data, columns=['T', 'F', 'L']).dropna().sort_values('T')
        sufijo = f"{dt_inicio.strftime('%y%m%d_%H%M')}_a_{dt_fin.strftime('%H%M')}"
        prefijo = fuente['nombre']

        print("💾 Generando entregables (Excel, CSV, Gráficas) por segmentos temporales...")
        sys.stdout.flush()

        tramos = []
        if ent.get('dividir', True):
            h_div = int(ent.get('horas_div', 1))
            t_curr = dt_inicio
            while t_curr < dt_fin:
                t_next = min(t_curr + datetime.timedelta(hours=h_div), dt_fin)
                tramos.append((t_curr, t_next))
                t_curr = t_next
        else:
            tramos.append((dt_inicio, dt_fin))

        datos_reporte_word_por_dia = {}
        dias_es = ["Lunes", "Martes", "Miercoles", "Jueves", "Viernes", "Sabado", "Domingo"]
        hay_alertas_global = False

        for tj, (t_ini, t_fin_tramo) in enumerate(tramos):
            df_tramo = df[(df['T'] >= t_ini) & (df['T'] <= t_fin_tramo)]
            if df_tramo.empty: continue
            
            nombre_dia = dias_es[t_ini.weekday()]
            carpeta_dia = f"{t_ini.strftime('%y%m%d')}_{nombre_dia}"
            ruta_dia = os.path.join(ruta_salida, carpeta_dia)
            os.makedirs(ruta_dia, exist_ok=True)
            
            suf_tramo = f"{t_ini.strftime('%y%m%d')}_{t_ini.strftime('%H%M')}_{t_fin_tramo.strftime('%H%M')}"
            tit_tramo = f"{t_ini.strftime('%d/%m/%Y %H:%M')} a {t_fin_tramo.strftime('%d/%m/%Y %H:%M')}"

            # --- EXPORTAR CSV Y EXCEL POR TRAMO Y POR DIA ---
            if ent['csv'] or ent['excel']:
                df_export = df_tramo.copy()
                df_export['F'] = (df_export['F'] * 1e6).astype(int)
                df_export['L'] = df_export['L'].round(1)
                
                if ent['csv']:
                    csv_path = os.path.join(ruta_dia, f"CSV_{prefijo}_{suf_tramo}.csv")
                    # Formato europeo requerido por el explorador
                    df_export['T'] = df_export['T'].dt.strftime('%d/%m/%Y %H:%M:%S.%f').str[:-3].str.replace('.', ',')
                    df_export.to_csv(csv_path, index=False, decimal=',', sep=';', header=["Tiempo", "Frecuencia (Hz)", "Level"])
                if ent['excel']:
                    excel_path = os.path.join(ruta_dia, f"EXCEL_{prefijo}_{suf_tramo}.xlsx")
                    try:
                        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                            df_export.iloc[:1048575].to_excel(writer, sheet_name='Hoja1', index=False)
                    except Exception as e: print(f"⚠️ Error EXCEL: {e}")

            # Evaluar alertas localmente
            df_maxhold, df_picos, p5_ruido, umbral, FI, FF, Central = calcular_metricas(df_tramo, OFFSET_ALERTAS, MAX_PICOS)
            hay_alertas = not df_picos.empty
            if hay_alertas: hay_alertas_global = True
            
            bloque_imgs = []

            if ent['graf_2dt'] or ent['graf_2df']:
                if ent['graf_2dt']:
                    fig_t = Figure(figsize=(15, 7), dpi=150); fig_t.patch.set_facecolor('white')
                    gs_t = gridspec.GridSpec(1, 2, width_ratios=[3.5, 1], wspace=0.05)
                    ax_t = fig_t.add_subplot(gs_t[0]); ax_side_t = fig_t.add_subplot(gs_t[1])
                    ax_t.plot(df_tramo['T'], df_tramo['L'], color='#0055A4', linewidth=0.5, alpha=0.8)
                    ax_t.axhline(y=umbral, color='orange', linestyle='--', linewidth=1.5, alpha=0.9)
                    if not df_picos.empty:
                        ax_t.scatter(df_picos['T'], df_picos['L'], color='red', s=30, zorder=5)
                        for _, r in df_picos.iterrows(): ax_t.annotate(str(int(r['Numero'])), (r['T'], r['L']), xytext=(0,6), textcoords="offset points", ha='center', fontsize=8, fontweight='bold', bbox=dict(boxstyle="circle,pad=0.2", fc="white", ec="red", lw=0.5))
                    ax_t.xaxis.set_major_formatter(mdates.DateFormatter('%H:%M:%S'))
                    ax_t.set_title(f"Evolución Temporal - {prefijo} ({tit_tramo})", fontsize=14, fontweight='bold')
                    ax_t.grid(True, linestyle=':', alpha=0.6)
                    dibujar_panel_lateral(ax_side_t, umbral, df_picos, OFFSET_ALERTAS)
                    ruta_img_t = os.path.join(ruta_dia, f"2D_TIME_{prefijo}_{suf_tramo}.png")
                    fig_t.savefig(ruta_img_t, bbox_inches='tight')
                    if hay_alertas: bloque_imgs.append(ruta_img_t)

                if ent['graf_2df']:
                    fig_f = Figure(figsize=(15, 7), dpi=150); fig_f.patch.set_facecolor('white')
                    gs_f = gridspec.GridSpec(1, 2, width_ratios=[3.5, 1], wspace=0.05)
                    ax_f = fig_f.add_subplot(gs_f[0]); ax_side_f = fig_f.add_subplot(gs_f[1])
                    ax_f.plot(df_maxhold['F'], df_maxhold['L'], color='#0055A4', linewidth=1.2)
                    ax_f.axhline(y=umbral, color='orange', linestyle='--', linewidth=1.5, alpha=0.9)
                    if not df_picos.empty:
                        ax_f.scatter(df_picos['F'], df_picos['L'], color='red', s=30, zorder=5)
                        for _, r in df_picos.iterrows(): ax_f.annotate(str(int(r['Numero'])), (r['F'], r['L']), xytext=(0,6), textcoords="offset points", ha='center', fontsize=8, fontweight='bold', bbox=dict(boxstyle="circle,pad=0.2", fc="white", ec="red", lw=0.5))
                        ax_f.text(0.02, 0.96, f"FI: {FI:.4f} MHz\nFF: {FF:.4f} MHz\nCentral: {Central:.4f} MHz", transform=ax_f.transAxes, fontsize=10, fontfamily='monospace', bbox=dict(boxstyle='round,pad=0.5', facecolor='#F8F9F9', edgecolor='#333333'))
                    
                    ax_f.ticklabel_format(useOffset=False, style='plain') # Desactivar notación científica E
                    ax_f.set_title(f"Espectro Frecuencias - {prefijo} ({tit_tramo})", fontsize=14, fontweight='bold')
                    ax_f.grid(True, linestyle=':', alpha=0.6)
                    dibujar_panel_lateral(ax_side_f, umbral, df_picos, OFFSET_ALERTAS)
                    ruta_img_f = os.path.join(ruta_dia, f"2D_FREQ_{prefijo}_{suf_tramo}.png")
                    fig_f.savefig(ruta_img_f, bbox_inches='tight')
                    if hay_alertas: bloque_imgs.append(ruta_img_f)

            if ent['graf_3d']:
                df_tramo_3d = df_tramo.copy()
                df_tramo_3d['F_bin'] = (df_tramo_3d['F'] / 0.05).round() * 0.05
                grid = df_tramo_3d.pivot_table(index=pd.Grouper(key='T', freq='10s'), columns='F_bin', values='L', aggfunc='max')
                suelo = df_tramo_3d['L'].min()
                grid = grid.fillna(suelo)
                X, Y = np.meshgrid(grid.columns.values, mdates.date2num(grid.index))
                Z = grid.values
                fig_3d = Figure(figsize=(16, 9), dpi=150)
                ax_3d = fig_3d.add_subplot(111, projection='3d')
                surf = ax_3d.plot_surface(X, Y, Z, cmap='turbo', linewidth=0, antialiased=False, vmin=suelo, vmax=df_tramo_3d['L'].max())
                ax_3d.view_init(elev=35, azim=-50)
                ax_3d.yaxis.set_major_formatter(mdates.DateFormatter('%H:%M:%S'))
                ax_3d.xaxis.set_major_formatter(ticker.ScalarFormatter(useOffset=False)) # Desactivar notación científica en 3D
                ax_3d.set_title(f"Espectro 3D {prefijo} ({tit_tramo})", pad=20)
                fig_3d.colorbar(surf, ax=ax_3d, shrink=0.5, aspect=15, pad=0.1)
                ruta_img_3d = os.path.join(ruta_dia, f"3D_{prefijo}_{suf_tramo}.png")
                fig_3d.savefig(ruta_img_3d, bbox_inches='tight')
                if hay_alertas: bloque_imgs.append(ruta_img_3d)
                
            if carpeta_dia not in datos_reporte_word_por_dia:
                datos_reporte_word_por_dia[carpeta_dia] = []

            datos_reporte_word_por_dia[carpeta_dia].append({
                "ini": t_ini.strftime('%d/%m/%Y %H:%M'),
                "fin": t_fin_tramo.strftime('%H:%M'),
                "hay_alertas": hay_alertas,
                "imgs": bloque_imgs
            })

        if ent['word'] and HAS_DOCX and datos_reporte_word_por_dia:
            for carpeta_dia, tramos_dia in datos_reporte_word_por_dia.items():
                ruta_dia = os.path.join(ruta_salida, carpeta_dia)
                doc = Document()
                doc.add_heading(f'Reporte de Alertas - {prefijo} ({carpeta_dia})', 0)
                
                for tramo in tramos_dia:
                    doc.add_heading(f"Tramo: {tramo['ini']} a {tramo['fin']}", level=1)
                    if not tramo['hay_alertas']:
                        p = doc.add_paragraph("No hay alertas en este rango de horas.")
                        p.paragraph_format.space_after = Inches(0.2)
                    else:
                        for ruta_img in tramo['imgs']:
                            if os.path.exists(ruta_img):
                                try:
                                    doc.add_picture(ruta_img, width=Inches(6.5))
                                    doc.add_paragraph(" ")
                                except Exception as e_img: print(f"Aviso al añadir imagen a Word: {e_img}")
                        
                doc.save(os.path.join(ruta_dia, f"Reporte_Alertas_{prefijo}_{carpeta_dia}.docx"))

        # --- ENVÍO DE EMAIL ---
        plt.close('all') 
        if payload.get('enviar_email') and payload.get('emails'):
            print("📧 Preparando envío de email...")
            enviar_correo_reporte(payload['emails'], ruta_salida, prefijo, titulo_graficas, hay_alertas_global)

        # Si era por red, ahora movemos todo del temporal local al destino final remoto
        if es_red:
            print(f"📦 Transfiriendo datos por red al destino final: {ruta_salida_final} ...")
            sys.stdout.flush()
            conectar_red_windows(base_res, fuente.get('usr_red'), fuente.get('pwd_red'))
            
            # Crear reintentos robustos de copiado
            import time
            copiado_ok = False
            error_msg = ""
            for intento in range(3):
                try:
                    os.makedirs(ruta_salida_final, exist_ok=True)
                    shutil.copytree(ruta_salida, ruta_salida_final, dirs_exist_ok=True)
                    copiado_ok = True
                    break
                except Exception as e_copy:
                    error_msg = str(e_copy)
                    print(f"⚠️ Reintento {intento+1} fallido al copiar por red: {e_copy}")
                    time.sleep(2)
                    
            if copiado_ok:
                try:
                    shutil.rmtree(ruta_salida, ignore_errors=True)
                except: pass
                ruta_salida_return = ruta_salida_final
            else:
                print(f"❌ Error definitivo copiando a red. Datos guardados en local: {ruta_salida}")
                ruta_salida_return = f"LOCAL_POR_ERROR_RED: {ruta_salida}"
        else:
            ruta_salida_return = ruta_salida

        print(f"🎉 EXTRACCIÓN COMPLETADA CON ÉXITO: {ruta_salida_return}\n" + "="*60)
        sys.stdout.flush()
        return jsonify({"status": "ok", "ruta_salida": ruta_salida_return})
        
    except Exception as e:
        tb = traceback.format_exc()
        print("❌ ERROR CRÍTICO EN EXTRACCIÓN:")
        print(tb)
        sys.stdout.flush()
        return jsonify({"status": "error", "error": f"Excepción interna: {str(e)}"})

def abrir_navegador():
    time.sleep(1.5)
    webbrowser.open(f"http://127.0.0.1:{PUERTO}")

if __name__ == '__main__':
    print(f"Iniciando Servidor Web Argus (Waitress) en el puerto {PUERTO}...")
    threading.Thread(target=abrir_navegador, daemon=True).start()
    
    # Servidor de producción Waitress (más estable y multihilo)
    serve(app, host='0.0.0.0', port=PUERTO, threads=6)

