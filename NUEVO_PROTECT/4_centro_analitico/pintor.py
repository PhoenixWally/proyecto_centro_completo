import sys
import os
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
from matplotlib.figure import Figure
import matplotlib.dates as mdates
import matplotlib.ticker as ticker
import matplotlib.gridspec as gridspec
from mpl_toolkits.mplot3d import Axes3D

try:
    from docx import Document
    from docx.shared import Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def calcular_metricas(df, offset_db, max_picos=30):
    p5_ruido = np.percentile(df['L'], 5)
    umbral = p5_ruido + offset_db
    df_maxhold = df.groupby(df['F'].round(4))['L'].max().reset_index()

    df_over = df[df['L'] >= umbral].sort_values(by='T', ascending=True)
    picos_detectados = []

    for _, row in df_over.iterrows():
        es_nuevo = True
        for p in picos_detectados:
            if abs((row['T'] - p['T']).total_seconds()) < 60 and abs(row['F'] - p['F']) < 0.5:
                es_nuevo = False
                if row['L'] > p['L']:
                    p['T'] = row['T']; p['L'] = row['L']; p['F'] = row['F']
                break
        if es_nuevo:
            if len(picos_detectados) < max_picos:
                picos_detectados.append({'T': row['T'], 'F': row['F'], 'L': row['L']})

    df_picos = pd.DataFrame(picos_detectados)
    if not df_picos.empty:
        df_picos = df_picos.sort_values('T').reset_index(drop=True)
        df_picos['Numero'] = df_picos.index + 1
        return df_maxhold, df_picos, p5_ruido, umbral, df_picos['F'].min(), df_picos['F'].max(), (df_picos['F'].min() + df_picos['F'].max()) / 2
    return df_maxhold, df_picos, p5_ruido, umbral, 0.0, 0.0, 0.0

def dibujar_panel_lateral(ax_side, umbral, df_picos, offset_db):
    ax_side.axis('off'); ax_side.set_xlim(0, 1); ax_side.set_ylim(0, 1)
    ax_side.text(0.0, 0.98, "LEYENDA", fontsize=11, fontweight='bold', ha='left', va='top')
    ax_side.plot([0.0, 0.15], [0.93, 0.93], color='#0055A4', lw=1.5)
    ax_side.text(0.2, 0.93, "Nivel de Señal", fontsize=10, va='center')
    ax_side.plot([0.0, 0.15], [0.88, 0.88], color='orange', ls='--', lw=1.5)
    ax_side.text(0.2, 0.88, f"Umbral (+{offset_db}dB)", fontsize=10, va='center')
    ax_side.scatter([0.075], [0.83], color='red', s=40, zorder=5)
    ax_side.text(0.2, 0.83, "Detecciones", fontsize=10, va='center')
    ax_side.text(0.0, 0.75, f"TOP PICOS ({len(df_picos)})", fontsize=11, fontweight='bold', ha='left', va='top')

    if df_picos.empty: ax_side.text(0.0, 0.70, "Ningún pico detectado.", fontsize=9, color='gray')
    else:
        ax_side.text(0.0, 0.70, "Nº   Hora       Frec(MHz)   Nivel", fontsize=9, fontweight='bold', fontfamily='monospace')
        y_pos = 0.67; step = min(0.03, 0.65 / max(15, len(df_picos)))
        for _, row in df_picos.iterrows():
            ax_side.text(0.0, y_pos, f"{int(row['Numero']):02d} | {row['T'].strftime('%H:%M:%S')} | {row['F']:8.4f} | {row['L']:5.1f}", fontsize=8, fontfamily='monospace', va='center')
            y_pos -= step

if __name__ == "__main__":
    archivo_csv = sys.argv[1]
    carpeta_destino_base = sys.argv[2]
    prefijo = sys.argv[3]
    t_ini = pd.to_datetime(sys.argv[4])
    t_fin_tramo = pd.to_datetime(sys.argv[5])
    
    OFFSET_ALERTAS = 15.0 

    # Replicamos el guardado en subcarpetas por días que tenías originalmente
    dias_es = ["Lunes", "Martes", "Miercoles", "Jueves", "Viernes", "Sabado", "Domingo"]
    nombre_dia = dias_es[t_ini.weekday()]
    carpeta_dia = f"{t_ini.strftime('%y%m%d')}_{nombre_dia}"
    ruta_dia = os.path.join(carpeta_destino_base, carpeta_dia)
    os.makedirs(ruta_dia, exist_ok=True)

    suf_tramo = f"{t_ini.strftime('%y%m%d')}_{t_ini.strftime('%H%M')}_{t_fin_tramo.strftime('%H%M')}"
    tit_tramo = f"{t_ini.strftime('%d/%m/%Y %H:%M')} a {t_fin_tramo.strftime('%d/%m/%Y %H:%M')}"

    df_tramo = pd.read_csv(archivo_csv)
    df_tramo['T'] = pd.to_datetime(df_tramo['T'])

    df_maxhold, df_picos, p5_ruido, umbral, FI, FF, Central = calcular_metricas(df_tramo, OFFSET_ALERTAS)
    hay_alertas = not df_picos.empty
    bloque_imgs = []

    # --- Gráfica 2D Tiempo ---
    try:
        fig_t = Figure(figsize=(15, 7), dpi=150); fig_t.patch.set_facecolor('white')
        gs_t = gridspec.GridSpec(1, 2, width_ratios=[3.5, 1], wspace=0.05)
        ax_t = fig_t.add_subplot(gs_t[0]); ax_side_t = fig_t.add_subplot(gs_t[1])
        ax_t.plot(df_tramo['T'], df_tramo['L'], color='#0055A4', linewidth=0.5, alpha=0.8)
        ax_t.axhline(y=umbral, color='orange', linestyle='--', linewidth=1.5, alpha=0.9)
        if not df_picos.empty:
            ax_t.scatter(df_picos['T'], df_picos['L'], color='red', s=30, zorder=5)
            for _, r in df_picos.iterrows(): ax_t.annotate(str(int(r['Numero'])), (r['T'], r['L']), xytext=(0,6), textcoords="offset points", ha='center', fontsize=8, fontweight='bold', bbox=dict(boxstyle="circle,pad=0.2", fc="white", ec="red", lw=0.5))
        ax_t.xaxis.set_major_formatter(mdates.DateFormatter('%H:%M:%S'))
        ax_t.set_title(f"Evolución Temporal - {prefijo} ({tit_tramo})", fontsize=14, fontweight='bold')
        ax_t.grid(True, linestyle=':', alpha=0.6)
        dibujar_panel_lateral(ax_side_t, umbral, df_picos, OFFSET_ALERTAS)
        ruta_img_t = os.path.join(ruta_dia, f"2D_TIME_{prefijo}_{suf_tramo}.png")
        fig_t.savefig(ruta_img_t, bbox_inches='tight')
        if hay_alertas: bloque_imgs.append(ruta_img_t)
    except Exception as e:
        print(f"Error generando gráfica 2D Tiempo: {e}")

    # --- Gráfica 2D Frecuencia ---
    try:
        fig_f = Figure(figsize=(15, 7), dpi=150); fig_f.patch.set_facecolor('white')
        gs_f = gridspec.GridSpec(1, 2, width_ratios=[3.5, 1], wspace=0.05)
        ax_f = fig_f.add_subplot(gs_f[0]); ax_side_f = fig_f.add_subplot(gs_f[1])
        ax_f.plot(df_maxhold['F'], df_maxhold['L'], color='#0055A4', linewidth=1.2)
        ax_f.axhline(y=umbral, color='orange', linestyle='--', linewidth=1.5, alpha=0.9)
        if not df_picos.empty:
            ax_f.scatter(df_picos['F'], df_picos['L'], color='red', s=30, zorder=5)
            for _, r in df_picos.iterrows(): ax_f.annotate(str(int(r['Numero'])), (r['F'], r['L']), xytext=(0,6), textcoords="offset points", ha='center', fontsize=8, fontweight='bold', bbox=dict(boxstyle="circle,pad=0.2", fc="white", ec="red", lw=0.5))
            ax_f.text(0.02, 0.96, f"FI: {FI:.4f} MHz\nFF: {FF:.4f} MHz\nCentral: {Central:.4f} MHz", transform=ax_f.transAxes, fontsize=10, fontfamily='monospace', bbox=dict(boxstyle='round,pad=0.5', facecolor='#F8F9F9', edgecolor='#333333'))
        ax_f.ticklabel_format(useOffset=False, style='plain')
        ax_f.set_title(f"Espectro Frecuencias - {prefijo} ({tit_tramo})", fontsize=14, fontweight='bold')
        ax_f.grid(True, linestyle=':', alpha=0.6)
        dibujar_panel_lateral(ax_side_f, umbral, df_picos, OFFSET_ALERTAS)
        ruta_img_f = os.path.join(ruta_dia, f"2D_FREQ_{prefijo}_{suf_tramo}.png")
        fig_f.savefig(ruta_img_f, bbox_inches='tight')
        if hay_alertas: bloque_imgs.append(ruta_img_f)
    except Exception as e:
        print(f"Error generando gráfica 2D Frecuencia: {e}")

    # --- Gráfica 3D ---
    try:
        df_tramo_3d = df_tramo.copy()
        df_tramo_3d['F_bin'] = (df_tramo_3d['F'] / 0.05).round() * 0.05
        grid = df_tramo_3d.pivot_table(index=pd.Grouper(key='T', freq='10s'), columns='F_bin', values='L', aggfunc='max')
        suelo = df_tramo_3d['L'].min()
        grid = grid.fillna(suelo) # Rellenar huecos con el ruido base
        
        X, Y = np.meshgrid(grid.columns.values, mdates.date2num(grid.index))
        Z = grid.values
        fig_3d = Figure(figsize=(16, 9), dpi=150)
        ax_3d = fig_3d.add_subplot(111, projection='3d')
        surf = ax_3d.plot_surface(X, Y, Z, cmap='turbo', linewidth=0, antialiased=False, vmin=suelo, vmax=df_tramo_3d['L'].max())
        
        ax_3d.view_init(elev=35, azim=-50)
        ax_3d.yaxis.set_major_formatter(mdates.DateFormatter('%H:%M:%S'))
        ax_3d.xaxis.set_major_formatter(ticker.ScalarFormatter(useOffset=False))
        ax_3d.set_title(f"Espectro 3D {prefijo} ({tit_tramo})", pad=20)
        fig_3d.colorbar(surf, ax=ax_3d, shrink=0.5, aspect=15, pad=0.1)
        
        ruta_img_3d = os.path.join(ruta_dia, f"3D_{prefijo}_{suf_tramo}.png")
        fig_3d.savefig(ruta_img_3d, bbox_inches='tight')
        if hay_alertas: bloque_imgs.append(ruta_img_3d)
    except Exception as e:
        print(f"Error generando gráfica 3D: {e}")

    # --- Exportar CSV y Excel con todos los puntos (Formato original) ---
    df_export = df_tramo.copy()
    # Frecuencia en Hz
    df_export['F'] = (df_export['F'] * 1e6).astype(int)
    # Nivel redondeado a 1 decimal
    df_export['L'] = df_export['L'].round(1)

    ruta_csv = os.path.join(ruta_dia, f"CSV_{prefijo}_{suf_tramo}.csv")
    try:
        df_csv = df_export.copy()
        # Formato europeo requerido por el explorador
        df_csv['T'] = df_csv['T'].dt.strftime('%d/%m/%Y %H:%M:%S.%f').str[:-3].str.replace('.', ',')
        df_csv.to_csv(ruta_csv, index=False, decimal=',', sep=';', header=["Tiempo", "Frecuencia (Hz)", "Level"])
    except Exception as e:
        print(f"Aviso: No se pudo generar el CSV: {e}")

    ruta_excel = os.path.join(ruta_dia, f"EXCEL_{prefijo}_{suf_tramo}.xlsx")
    try:
        with pd.ExcelWriter(ruta_excel, engine='openpyxl') as writer:
            df_export.iloc[:1048575].to_excel(writer, sheet_name='Hoja1', index=False)
    except Exception as e:
        print(f"Aviso: No se pudo generar el Excel: {e}")
    # --- Generar Documento Word Automático ---
    if HAS_DOCX and hay_alertas:
        doc = Document()
        doc.add_heading(f'Reporte de Alertas - {prefijo} ({carpeta_dia})', 0)
        doc.add_heading(f"Tramo: {t_ini.strftime('%d/%m/%Y %H:%M')} a {t_fin_tramo.strftime('%H:%M')}", level=1)
        for ruta_img in bloque_imgs:
            if os.path.exists(ruta_img):
                doc.add_picture(ruta_img, width=Inches(6.5))
                doc.add_paragraph(" ")
        doc.save(os.path.join(ruta_dia, f"Reporte_Alertas_{prefijo}_{suf_tramo}.docx"))